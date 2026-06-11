from __future__ import annotations

from pathlib import Path
from textwrap import dedent

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = Path(r"D:\My\XAI\Doc\Bao_cao_hoc_phan_XAI_Medical_4_chuong.docx")
ASSET_DIR = Path(r"D:\My\XAI\Doc\report_assets_4chapters")
ASSET_DIR.mkdir(parents=True, exist_ok=True)


TITLE = "HỆ THỐNG XAI MEDICAL HỖ TRỢ CHẨN ĐOÁN TRÀN DỊCH MÀNG PHỔI TỪ ẢNH X-QUANG NGỰC"
COURSE = "TRÍ TUỆ NHÂN TẠO ỨNG DỤNG TRONG Y TẾ"
SCHOOL = "TRƯỜNG ĐẠI HỌC ĐẠI NAM"


def font(size: int, bold: bool = False):
    candidates = [
        r"C:\Windows\Fonts\times.ttf",
        r"C:\Windows\Fonts\timesbd.ttf" if bold else r"C:\Windows\Fonts\times.ttf",
        r"C:\Windows\Fonts\arial.ttf",
    ]
    for p in candidates:
        if p and Path(p).exists():
            return ImageFont.truetype(p, size)
    return ImageFont.load_default()


def make_canvas(name: str, title: str):
    img = Image.new("RGB", (1600, 900), "#FFFFFF")
    d = ImageDraw.Draw(img)
    d.rectangle([0, 0, 1600, 100], fill="#17365D")
    d.text((60, 30), title, fill="white", font=font(36, True))
    return img, d


def save_architecture():
    img, d = make_canvas("architecture", "Kiến trúc tổng thể hệ thống XAI Medical")
    boxes = [
        ("Người dùng\nUpload ảnh X-quang", 80, 230, "#D9EAF7"),
        ("Frontend React/Vite\nDashboard, Diagnosis, Batch", 430, 230, "#E2F0D9"),
        ("FastAPI Backend\n/predict, /predict_batch, /stats", 800, 230, "#FFF2CC"),
        ("AI Engine\nDenseNet-121 + Grad-CAM", 1160, 230, "#FCE4D6"),
    ]
    for text, x, y, color in boxes:
        d.rounded_rectangle([x, y, x + 300, y + 180], radius=24, fill=color, outline="#17365D", width=4)
        d.multiline_text((x + 24, y + 42), text, fill="#111111", font=font(28, True), spacing=8)
    for x1, x2 in [(380, 430), (730, 800), (1100, 1160)]:
        d.line([x1, 320, x2, 320], fill="#17365D", width=8)
        d.polygon([(x2, 320), (x2 - 24, 305), (x2 - 24, 335)], fill="#17365D")
    d.rounded_rectangle([360, 570, 1240, 730], radius=24, fill="#F2F2F2", outline="#7F7F7F", width=3)
    d.text((405, 610), "SQLite lưu lịch sử, html2canvas + jsPDF xuất báo cáo, OOD kiểm tra ảnh ngoài phân phối", fill="#111111", font=font(28))
    out = ASSET_DIR / "01_kien_truc_tong_the.png"
    img.save(out)
    return out


def save_dataset_chart():
    img, d = make_canvas("dataset", "Quy mô dữ liệu pretrained của TorchXRayVision weights all")
    data = [("NIH ChestX-ray14", 112120, "#5B9BD5"), ("CheXpert", 224316, "#70AD47"), ("MIMIC-CXR", 377110, "#FFC000"), ("PadChest", 160000, "#ED7D31")]
    maxv = max(v for _, v, _ in data)
    y = 210
    for label, val, color in data:
        w = int(900 * val / maxv)
        d.text((120, y + 18), label, fill="#111111", font=font(28, True))
        d.rounded_rectangle([450, y, 450 + w, y + 70], radius=14, fill=color)
        d.text((470 + w, y + 18), f"{val:,} ảnh".replace(",", "."), fill="#111111", font=font(26))
        y += 120
    d.text((120, 740), "Tổng quy mô xấp xỉ trên 800.000 ảnh X-quang ngực; hệ thống chỉ fine-tune classifier Effusion.", fill="#111111", font=font(28, True))
    out = ASSET_DIR / "02_bieu_do_dataset.png"
    img.save(out)
    return out


def save_pipeline():
    img, d = make_canvas("pipeline", "Pipeline tiền xử lý và suy luận")
    steps = [
        "Ảnh JPEG/PNG",
        "Kiểm tra OOD",
        "Grayscale",
        "XRayResizer(224)",
        "Normalize [-1024,1024]",
        "DenseNet-121",
        "Effusion probability",
        "PDF/History",
    ]
    x = 60
    for i, step in enumerate(steps):
        d.rounded_rectangle([x, 310, x + 170, 470], radius=18, fill="#DDEBF7", outline="#17365D", width=3)
        d.multiline_text((x + 16, 350), step, fill="#111111", font=font(23, True), spacing=4)
        if i < len(steps) - 1:
            d.line([x + 175, 390, x + 215, 390], fill="#17365D", width=6)
            d.polygon([(x + 215, 390), (x + 198, 378), (x + 198, 402)], fill="#17365D")
        x += 190
    d.text((80, 610), "Điểm khác biệt quan trọng: hệ thống không dùng preprocessing ImageNet; toàn bộ ảnh được đưa về chuẩn X-quang của TorchXRayVision.", fill="#111111", font=font(30))
    out = ASSET_DIR / "03_pipeline_tien_xu_ly.png"
    img.save(out)
    return out


def save_gradcam():
    img, d = make_canvas("gradcam", "Nguyên lý Grad-CAM trong mô hình phân loại")
    d.rounded_rectangle([100, 210, 460, 640], radius=28, fill="#F2F2F2", outline="#17365D", width=4)
    d.ellipse([190, 300, 285, 540], fill="#A6A6A6")
    d.ellipse([285, 300, 380, 540], fill="#A6A6A6")
    d.text((150, 670), "Ảnh X-quang đầu vào", fill="#111111", font=font(28, True))
    d.rounded_rectangle([620, 250, 960, 590], radius=28, fill="#FFF2CC", outline="#17365D", width=4)
    d.text((670, 380), "denseblock4\n1024 kênh\nfeature maps", fill="#111111", font=font(34, True), spacing=10)
    d.rounded_rectangle([1120, 210, 1480, 640], radius=28, fill="#F2F2F2", outline="#17365D", width=4)
    d.ellipse([1210, 300, 1305, 540], fill="#A6A6A6")
    d.ellipse([1305, 300, 1400, 540], fill="#A6A6A6")
    d.ellipse([1330, 460, 1440, 570], fill="#FF0000")
    d.ellipse([1365, 495, 1415, 545], fill="#FFC000")
    d.text((1180, 670), "Heatmap giải thích", fill="#111111", font=font(28, True))
    for x1, x2 in [(470, 620), (960, 1120)]:
        d.line([x1, 420, x2, 420], fill="#17365D", width=8)
        d.polygon([(x2, 420), (x2 - 28, 402), (x2 - 28, 438)], fill="#17365D")
    out = ASSET_DIR / "04_gradcam_denseblock4.png"
    img.save(out)
    return out


def save_threshold():
    img, d = make_canvas("threshold", "Ngưỡng vận hành theo Youden's J statistic")
    d.line([180, 720, 1420, 720], fill="#111111", width=4)
    d.line([180, 720, 180, 170], fill="#111111", width=4)
    points = [(180, 700), (360, 610), (560, 470), (760, 350), (980, 270), (1220, 210), (1420, 185)]
    d.line(points, fill="#5B9BD5", width=8)
    d.line([760, 720, 760, 350], fill="#C00000", width=5)
    d.ellipse([742, 332, 778, 368], fill="#C00000")
    d.text((800, 360), "Optimal threshold = 0,0682", fill="#C00000", font=font(34, True))
    d.text((650, 765), "False Positive Rate", fill="#111111", font=font(28))
    d.text((40, 145), "True Positive Rate", fill="#111111", font=font(28))
    d.text((220, 220), "AUC = 0,87; Accuracy = 81,46% trên hold-out test set Cloud/Kaggle", fill="#111111", font=font(30, True))
    out = ASSET_DIR / "05_roc_threshold.png"
    img.save(out)
    return out


def save_pdf_flow():
    img, d = make_canvas("pdf", "Luồng xuất báo cáo PDF và lưu lịch sử")
    items = [
        ("Result Modal", 90, 320, "#D9EAF7"),
        ("html2canvas", 420, 320, "#E2F0D9"),
        ("jsPDF", 750, 320, "#FFF2CC"),
        ("PDF tải về", 1080, 320, "#FCE4D6"),
    ]
    for label, x, y, color in items:
        d.rounded_rectangle([x, y, x + 250, y + 150], radius=22, fill=color, outline="#17365D", width=4)
        d.text((x + 36, y + 56), label, fill="#111111", font=font(32, True))
    for x1, x2 in [(345, 420), (675, 750), (1005, 1080)]:
        d.line([x1, 395, x2, 395], fill="#17365D", width=8)
        d.polygon([(x2, 395), (x2 - 24, 378), (x2 - 24, 412)], fill="#17365D")
    d.rounded_rectangle([260, 590, 1340, 710], radius=18, fill="#F2F2F2", outline="#7F7F7F", width=3)
    d.text((300, 630), "Cách xuất mới giữ nguyên bố cục tiếng Việt, ảnh gốc, heatmap, confidence và giải thích y khoa.", fill="#111111", font=font(30))
    out = ASSET_DIR / "06_luong_xuat_pdf.png"
    img.save(out)
    return out


def create_images():
    return [
        save_architecture(),
        save_dataset_chart(),
        save_pipeline(),
        save_gradcam(),
        save_threshold(),
        save_pdf_flow(),
    ]


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) < 35 else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(12)
    r.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True)
        set_cell_shading(table.rows[0].cells[i], "D9EAF7")
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            set_cell_text(row.cells[i], val)
    if widths:
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    return table


def set_document_defaults(doc: Document):
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Cm(1.0)
    normal.paragraph_format.space_after = Pt(6)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        st = styles[name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.line_spacing = 1.5
        st.paragraph_format.space_before = Pt(6)
        st.paragraph_format.space_after = Pt(6)
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.bold = True
    styles["Heading 3"].font.size = Pt(13)
    styles["Heading 3"].font.bold = True
    styles["Heading 3"].font.italic = True


def add_para(doc: Document, text: str, italic: bool = False, bold: bool = False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text.strip())
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(13)
    r.italic = italic
    r.bold = bold
    return p


def add_chapter(doc: Document, title: str):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title.upper())
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(14)
    r.bold = True
    return p


def add_h2(doc: Document, text: str):
    p = doc.add_heading(text, level=2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_h3(doc: Document, text: str):
    p = doc.add_heading(text, level=3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(12)
    r.italic = True


def add_picture(doc: Document, path: Path, caption: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(5.9))
    add_caption(doc, caption)


def cover(doc: Document):
    for text in ["BỘ GIÁO DỤC VÀ ĐÀO TẠO", SCHOOL]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r.font.size = Pt(13)
        r.bold = True
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BÁO CÁO HỌC PHẦN")
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(18)
    r.bold = True
    for line in [f"TÊN HỌC PHẦN: {COURSE}", f"ĐỀ TÀI: {TITLE}"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        r = p.add_run(line)
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r.font.size = Pt(14)
        r.bold = True
    doc.add_paragraph()
    add_table(doc, ["STT", "Mã sinh viên", "Họ và tên", "Ngày sinh", "Lớp"], [["1", "", "", "", ""]], [1.1, 3.0, 5.0, 3.0, 3.0])
    for _ in range(7):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run("Hà Nội, năm 2026")
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(13)
    doc.add_page_break()


def front_matter(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("LỜI NÓI ĐẦU")
    r.bold = True
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(14)
    add_para(doc, "Báo cáo học phần này được xây dựng nhằm trình bày một cách hệ thống quá trình nghiên cứu, thiết kế, triển khai và đánh giá sản phẩm XAI Medical, một hệ thống hỗ trợ chẩn đoán tràn dịch màng phổi từ ảnh X-quang ngực. Nội dung báo cáo tập trung vào dữ liệu huấn luyện, mô hình DenseNet-121/TorchXRayVision, cơ chế fine-tune cho nhãn Effusion, kỹ thuật giải thích Grad-CAM, kiến trúc phần mềm FastAPI và React, đồng thời phân tích các quyết định kỹ thuật đã được lựa chọn trong quá trình phát triển sản phẩm. Do đây là một hệ thống liên quan đến y tế, báo cáo luôn nhấn mạnh giới hạn sử dụng: kết quả AI chỉ đóng vai trò hỗ trợ tham khảo và không thay thế bác sĩ chuyên khoa chẩn đoán hình ảnh.")
    add_para(doc, "Báo cáo được viết theo hướng hạn chế tách dòng nhỏ lẻ, ưu tiên trình bày bằng các đoạn văn hoàn chỉnh để bảo đảm mạch lập luận rõ ràng, đồng thời vẫn sử dụng hình ảnh, bảng biểu và sơ đồ tại những vị trí cần minh họa. Các số liệu chính trong báo cáo được dùng theo thông tin dự án đã xác nhận: mô hình gốc sử dụng weights `densenet121-res224-all`, được huấn luyện trên NIH ChestX-ray14, CheXpert, MIMIC-CXR và PadChest; mô hình fine-tune cho Effusion dùng trọng số `pleural_effusion_model_best.pth`; kết quả ghi nhận gồm Accuracy 81,46%, AUC 0,87 và ngưỡng vận hành 0,0682 theo Youden's J statistic.")
    doc.add_page_break()
    for title, entries in [
        ("MỤC LỤC", ["CHƯƠNG 1. TỔNG QUAN VÀ CƠ SỞ LÝ THUYẾT", "CHƯƠNG 2. DỮ LIỆU, MÔ HÌNH HUẤN LUYỆN VÀ PHƯƠNG PHÁP", "CHƯƠNG 3. THIẾT KẾ VÀ XÂY DỰNG HỆ THỐNG", "CHƯƠNG 4. THỰC NGHIỆM, ĐÁNH GIÁ VÀ HOÀN THIỆN SẢN PHẨM", "KẾT LUẬN", "DANH MỤC TÀI LIỆU THAM KHẢO"]),
        ("MỤC LỤC HÌNH ẢNH", ["Hình 1.1. Kiến trúc tổng thể hệ thống XAI Medical", "Hình 2.1. Quy mô dữ liệu pretrained", "Hình 2.2. Pipeline tiền xử lý và suy luận", "Hình 2.3. Nguyên lý Grad-CAM tại denseblock4", "Hình 4.1. Ngưỡng vận hành theo ROC", "Hình 4.2. Luồng xuất báo cáo PDF"]),
        ("MỤC LỤC BẢNG", ["Bảng 1.1. Bảng từ viết tắt", "Bảng 2.1. Nguồn dữ liệu và quy mô", "Bảng 2.2. Nhãn bệnh lý đa nhãn của mô hình", "Bảng 3.1. Thành phần kiến trúc hệ thống", "Bảng 4.1. Thông số đánh giá và triển khai"]),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        r = p.add_run(title)
        r.bold = True
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r.font.size = Pt(14)
        for e in entries:
            add_para(doc, e)
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run("BẢNG CÁC TỪ VIẾT TẮT")
    r.bold = True
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(14)
    add_table(doc, ["STT", "Từ viết tắt", "Viết đầy đủ"], [
        ["1", "AI", "Artificial Intelligence - Trí tuệ nhân tạo"],
        ["2", "XAI", "Explainable Artificial Intelligence - Trí tuệ nhân tạo giải thích được"],
        ["3", "CNN", "Convolutional Neural Network - Mạng nơ-ron tích chập"],
        ["4", "CXR", "Chest X-ray - Ảnh X-quang ngực"],
        ["5", "Grad-CAM", "Gradient-weighted Class Activation Mapping"],
        ["6", "AUC", "Area Under the ROC Curve"],
        ["7", "OOD", "Out-of-Distribution - Ảnh ngoài phân phối"],
        ["8", "API", "Application Programming Interface"],
        ["9", "DICOM", "Digital Imaging and Communications in Medicine"],
    ], [1.2, 3.2, 11.0])


SECTION_PARAS = {
    "1.1. Lý do chọn đề tài": [
        "Ảnh X-quang ngực là một trong những phương tiện chẩn đoán hình ảnh phổ biến nhất trong thực hành lâm sàng vì chi phí thấp, tốc độ chụp nhanh và khả năng cung cấp thông tin tổng quan về phổi, màng phổi, tim và lồng ngực. Trong đó, tràn dịch màng phổi là tình trạng tích tụ dịch bất thường trong khoang màng phổi, có thể liên quan đến nhiều nguyên nhân như suy tim, viêm phổi, lao, ung thư, chấn thương hoặc bệnh lý viêm hệ thống. Trên phim X-quang, dấu hiệu có thể biểu hiện bằng mờ góc sườn hoành, tăng cản quang vùng đáy phổi, đường cong Damoiseau hoặc các thay đổi hình thái khác tùy mức độ dịch. Việc đọc phim đòi hỏi kinh nghiệm chuyên môn, đồng thời dễ bị ảnh hưởng bởi chất lượng ảnh, tư thế chụp, bệnh lý phối hợp và áp lực thời gian trong môi trường bệnh viện.",
        "Trong bối cảnh trí tuệ nhân tạo phát triển mạnh, các mô hình học sâu có thể hỗ trợ sàng lọc nhanh ảnh X-quang, đặc biệt khi số lượng ảnh lớn hoặc khi cần ưu tiên ca nghi ngờ. Tuy nhiên, hệ thống AI y tế không thể chỉ đưa ra một nhãn dự đoán và một con số xác suất, bởi bác sĩ cần biết vì sao hệ thống đưa ra kết luận đó. Đây là lý do XAI trở thành yêu cầu quan trọng: kết quả dự đoán phải đi kèm bằng chứng trực quan, ví dụ heatmap thể hiện vùng ảnh có ảnh hưởng mạnh đến quyết định của mô hình. Đề tài XAI Medical được lựa chọn vì kết hợp được ba lớp vấn đề: bài toán y khoa có ý nghĩa thực tiễn, mô hình học sâu chuyên biệt miền X-quang và hệ thống phần mềm có thể triển khai thành sản phẩm sử dụng được.",
        "Một điểm cần nhấn mạnh là hệ thống trong báo cáo không được mô tả như một công cụ thay thế bác sĩ. Vai trò đúng của sản phẩm là hỗ trợ ra quyết định, cung cấp thêm góc nhìn định lượng và trực quan để người dùng tham khảo. Trong y tế, sai số mô hình có thể gây hậu quả nghiêm trọng nếu được hiểu như kết luận cuối cùng, vì vậy báo cáo luôn đặt mô hình trong giới hạn của một hệ thống hỗ trợ. Cách tiếp cận này phù hợp với định hướng xây dựng sản phẩm có trách nhiệm: dùng AI để tăng tốc phân tích và chuẩn hóa một phần quy trình, nhưng vẫn giữ quyền quyết định cuối cùng cho chuyên gia lâm sàng.",
    ],
    "1.2. Mục tiêu và phạm vi nghiên cứu": [
        "Mục tiêu tổng quát của báo cáo là trình bày quá trình xây dựng hệ thống XAI Medical có khả năng nhận ảnh X-quang ngực, kiểm tra tính phù hợp của ảnh đầu vào, tiền xử lý theo chuẩn TorchXRayVision, suy luận bằng mô hình DenseNet-121, trích xuất xác suất nhãn Effusion, tạo heatmap Grad-CAM và hiển thị kết quả trên giao diện web. Báo cáo không chỉ dừng ở mô hình mà còn mô tả toàn bộ pipeline sản phẩm, từ frontend, backend, lưu trữ SQLite, xử lý hàng loạt, cơ chế giải thích, kiểm soát ảnh ngoài phân phối cho đến chức năng xuất báo cáo PDF.",
        "Phạm vi nghiên cứu tập trung vào bài toán tràn dịch màng phổi, tức nhãn Effusion trong danh sách pathology của mô hình đa nhãn. Mô hình nền không được huấn luyện từ đầu trong phạm vi báo cáo, mà kế thừa pretrained DenseNet-121 của TorchXRayVision với weights `densenet121-res224-all`, sau đó nạp trọng số fine-tune `pleural_effusion_model_best.pth` để tập trung vào classifier Effusion. Việc trình bày như vậy giúp phân biệt rõ giữa năng lực học đặc trưng tổng quát từ hơn 800.000 ảnh X-quang và bước tinh chỉnh phục vụ mục tiêu sản phẩm cụ thể.",
        "Báo cáo cũng đặt ra các giới hạn rõ ràng để tránh hiểu sai. Grad-CAM trong hệ thống là kỹ thuật giải thích của mô hình phân loại, không phải kỹ thuật phân đoạn chính xác vùng bệnh. Do đó heatmap phản ánh vùng ảnh làm tăng kích hoạt mô hình, chứ không phải đường biên tổn thương đã được bác sĩ đánh dấu. Các thuật toán hậu xử lý như contour, alpha blending hoặc kiểm soát màu sắc chỉ nhằm cải thiện khả năng đọc và thẩm mỹ, không làm thay đổi trọng số mô hình đã huấn luyện. Phạm vi này bảo đảm sản phẩm giữ nguyên dữ liệu và mô hình đã huấn luyện, phù hợp yêu cầu không huấn luyện thêm AI khác.",
    ],
    "1.3. Cơ sở lý thuyết về CNN và DenseNet": [
        "Mạng nơ-ron tích chập là kiến trúc học sâu đặc biệt phù hợp với dữ liệu ảnh vì các lớp tích chập có khả năng học đặc trưng cục bộ như cạnh, vùng sáng tối, kết cấu mô và dần tổng hợp thành đặc trưng mức cao hơn. Đối với ảnh X-quang, các đặc trưng quan trọng thường rất tinh tế, ví dụ sự mờ nhẹ ở đáy phổi, thay đổi biên cơ hoành hoặc phân bố cản quang không đồng nhất. CNN học các đặc trưng này thông qua nhiều tầng lọc, sau đó lớp phân loại chuyển biểu diễn ảnh thành xác suất bệnh lý. Khả năng học tự động đặc trưng giúp CNN vượt qua hạn chế của các phương pháp thủ công vốn phụ thuộc nhiều vào lựa chọn đặc trưng do con người thiết kế.",
        "DenseNet là một biến thể CNN nổi bật vì sử dụng kết nối dày đặc giữa các tầng, nghĩa là mỗi tầng nhận đầu vào từ nhiều tầng trước đó thay vì chỉ nhận từ tầng ngay trước. Cơ chế này giúp cải thiện dòng gradient, tái sử dụng đặc trưng và giảm nguy cơ mất thông tin trong mạng sâu. Với ảnh y tế, việc giữ lại cả đặc trưng mức thấp và mức cao là quan trọng vì tổn thương có thể biểu hiện bằng thay đổi hình thái nhỏ. DenseNet-121 trở thành lựa chọn phù hợp cho sản phẩm vì cân bằng giữa độ sâu, hiệu quả tính toán và mức độ phổ biến trong các nghiên cứu X-quang ngực.",
        "Trong hệ thống XAI Medical, DenseNet-121 không được dùng như một mô hình ImageNet thông thường mà được lấy từ TorchXRayVision, một thư viện chuyên cho ảnh X-quang ngực. Sự khác biệt này quan trọng vì ảnh X-quang có phân phối pixel, cấu trúc giải phẫu và ý nghĩa đặc trưng khác ảnh tự nhiên. Nếu dùng mô hình tiền huấn luyện trên ảnh tự nhiên mà không điều chỉnh, hệ thống có thể học các đặc trưng không phù hợp. Việc dùng TorchXRayVision giúp pipeline preprocessing, danh sách pathology và trọng số pretrained thống nhất trong cùng miền dữ liệu y khoa.",
    ],
    "1.4. Cơ sở lý thuyết về XAI và Grad-CAM": [
        "XAI là hướng nghiên cứu nhằm làm cho mô hình AI có thể giải thích, kiểm chứng hoặc ít nhất là quan sát được ở mức quyết định. Trong chẩn đoán hình ảnh, yêu cầu này càng quan trọng vì người dùng không thể chỉ dựa vào một điểm số để ra quyết định y khoa. Grad-CAM là một kỹ thuật phổ biến để trực quan hóa vùng ảnh ảnh hưởng đến dự đoán của CNN bằng cách sử dụng gradient của lớp mục tiêu đối với feature maps ở tầng sâu. Kết quả là một bản đồ nhiệt, trong đó vùng nóng cho biết các khu vực có đóng góp lớn hơn vào dự đoán.",
        "Điểm cần hiểu đúng là Grad-CAM không tạo mask tổn thương theo nghĩa phân đoạn y khoa. Nó cho biết mô hình đã chú ý vào đâu, không bảo đảm vùng chú ý trùng khít hoàn toàn với ranh giới bệnh lý. Đối với tràn dịch màng phổi, mô hình có thể tập trung vào góc sườn hoành, đáy phổi hoặc vùng mờ cản quang, nhưng nếu feature map có độ phân giải thấp thì heatmap khi phóng to có thể bị lem. Vì vậy báo cáo không mô tả heatmap như nhãn vàng, mà xem đây là bằng chứng trực quan hỗ trợ diễn giải mô hình.",
        "Trong quá trình triển khai, hệ thống từng cân nhắc hook vào tầng `norm5`, nhưng do đặc điểm ReLU inplace trong TorchXRayVision có thể gây lỗi autograd khi tính gradient, hệ thống chọn `denseblock4` làm tầng mục tiêu. Lựa chọn này an toàn hơn về kỹ thuật, giữ được 1024 kênh đặc trưng sâu và vẫn cung cấp thông tin đủ mạnh cho Grad-CAM. Đây là một ví dụ cho thấy sản phẩm không chỉ áp dụng lý thuyết mà còn cần xử lý các vấn đề thực tế của framework học sâu.",
    ],
    "1.5. Tổng quan sản phẩm XAI Medical": [
        "XAI Medical được thiết kế như một hệ thống web có đầy đủ các thành phần của một sản phẩm thử nghiệm gần mức triển khai: giao diện người dùng React/Vite, backend FastAPI, lõi suy luận PyTorch/TorchXRayVision, cơ chế lưu lịch sử SQLite, dashboard thống kê, xử lý ảnh đơn lẻ và hàng loạt, modal kết quả, heatmap giải thích và xuất báo cáo PDF. Cách tiếp cận này khác với nhiều bài thực nghiệm mô hình thuần túy vì người dùng có thể tương tác trực tiếp với sản phẩm, tải ảnh lên, xem kết quả và lưu lại báo cáo.",
        "Về mặt nghiệp vụ, hệ thống nhận ảnh X-quang ngực, kiểm tra xem ảnh có phù hợp miền dữ liệu hay không, chuẩn hóa ảnh về grayscale, resize về 224 x 224, normalize theo chuẩn TorchXRayVision, suy luận xác suất Effusion và so sánh với ngưỡng 0,0682. Nếu xác suất vượt ngưỡng, hệ thống gán nhãn tràn dịch màng phổi; nếu không, hệ thống gán nhãn bình thường theo phạm vi bài toán Effusion. Độ tin cậy hiển thị cũng được xử lý đúng theo nhãn: nếu có bệnh thì dùng xác suất bệnh, nếu bình thường thì dùng 1 trừ xác suất bệnh để tránh hiểu nhầm.",
        "Tổng thể sản phẩm thể hiện một workflow khép kín: mô hình học sâu đưa ra dự đoán, Grad-CAM giải thích vùng ảnh liên quan, frontend trình bày kết quả bằng giao diện trực quan, backend lưu thông tin bệnh án và PDF giúp xuất kết quả thành tài liệu. Đây là cơ sở để mở rộng sản phẩm trong tương lai, ví dụ hỗ trợ DICOM, phân quyền người dùng, triển khai Docker hoặc bổ sung báo cáo định lượng đầy đủ hơn như confusion matrix và ROC artifact.",
    ],
}


CH2_PARAS = {
    "2.1. Nguồn dữ liệu pretrained": [
        "Mô hình gốc trong hệ thống sử dụng weights `densenet121-res224-all`, trong đó `all` thể hiện việc mô hình được huấn luyện trên tổ hợp nhiều bộ dữ liệu X-quang ngực công khai thay vì một nguồn duy nhất. Bốn nguồn chính gồm NIH ChestX-ray14, CheXpert, MIMIC-CXR và PadChest. NIH ChestX-ray14 cung cấp 112.120 ảnh và là nguồn gốc của con số thường được nhắc đến trong các tài liệu DenseNet X-quang; CheXpert đóng góp khoảng 224.316 ảnh; MIMIC-CXR đóng góp khoảng 377.110 ảnh; PadChest đóng góp khoảng 160.000 ảnh. Tổng quy mô xấp xỉ hơn 800.000 ảnh, tạo nền tảng học đặc trưng phong phú cho mô hình pretrained.",
        "Việc gộp nhiều dataset có ý nghĩa quan trọng vì ảnh X-quang trong thực tế thay đổi theo bệnh viện, máy chụp, tư thế, chất lượng phim, đặc điểm dân số và quy trình gán nhãn. Một mô hình chỉ học từ một nguồn có thể phụ thuộc vào đặc trưng riêng của nguồn đó, trong khi mô hình học từ nhiều nguồn có khả năng khái quát tốt hơn. Tuy vậy, điều này không đồng nghĩa mô hình hoàn hảo trong mọi bối cảnh; khi triển khai thực tế vẫn cần đánh giá trên tập test nội bộ hoặc dữ liệu của cơ sở y tế mục tiêu.",
        "Trong báo cáo, cần phân biệt rõ dataset pretrained và dataset kiểm thử thực tế. Bộ 50 ảnh COVID-19/Pneumonia tải từ repo `ieee8023/covid-chestxray-dataset` chỉ dùng để kiểm thử hành vi hệ thống và giao diện, không phải tập tạo ra Accuracy 81,46% và AUC 0,87. Hai chỉ số này thuộc hold-out test set trong quá trình fine-tuning Cloud/Kaggle. Phân biệt này giúp tránh sai lệch khi trình bày kết quả, đặc biệt trong báo cáo học thuật.",
    ],
    "2.2. Tiền xử lý ảnh X-quang": [
        "Pipeline tiền xử lý được thiết kế nghiêm ngặt theo miền X-quang. Ảnh đầu vào từ người dùng có thể ở dạng JPEG hoặc PNG, nhưng hệ thống kiểm tra định dạng, kích thước và các dấu hiệu ngoài phân phối trước khi đưa vào mô hình. Ảnh không phải X-quang, ví dụ ảnh phong cảnh hoặc ảnh màu tự nhiên, không nên được đưa trực tiếp vào mô hình vì mô hình chỉ học phân phối ảnh X-quang. Bước kiểm soát OOD giúp giảm rủi ro hệ thống trả kết quả có vẻ hợp lệ cho dữ liệu vô nghĩa.",
        "Sau khi ảnh được chấp nhận, hệ thống chuyển ảnh về grayscale. Đây là yêu cầu phù hợp với bản chất ảnh X-quang, nơi thông tin chính nằm trong cường độ sáng tối chứ không phải màu sắc. Ảnh tiếp tục được resize bằng `XRayResizer(224)` về kích thước 224 x 224 để tương thích với backbone DenseNet-121 của TorchXRayVision. Kích thước này không phải tùy chọn ngẫu nhiên mà gắn với cấu hình mô hình pretrained, bảo đảm tensor đầu vào đúng hình dạng `(1, 1, 224, 224)`.",
        "Bước chuẩn hóa sử dụng `xrv.datasets.normalize(img_np, 255)`, đưa pixel từ miền [0, 255] về chuẩn X-quang khoảng [-1024, 1024]. Đây là khác biệt rất quan trọng so với các mô hình ảnh tự nhiên thường dùng mean/std của ImageNet. Nếu dùng chuẩn hóa sai miền, feature map và xác suất đầu ra có thể bị lệch vì mô hình được huấn luyện trên phân phối khác. Do đó preprocessing là một phần không thể tách rời của mô hình, không chỉ là thao tác phụ trợ.",
    ],
    "2.3. Mô hình đa nhãn và nhãn Effusion": [
        "DenseNet-121 trong TorchXRayVision là mô hình phân loại đa nhãn, nghĩa là một ảnh có thể đồng thời mang nhiều nhãn bệnh lý. Danh sách pathology bao gồm các nhãn như Atelectasis, Consolidation, Infiltration, Pneumothorax, Edema, Emphysema, Fibrosis, Effusion, Pneumonia, Pleural Thickening, Cardiomegaly, Nodule, Mass và Hernia. Cách học đa nhãn phù hợp với ảnh X-quang vì một bệnh nhân có thể có nhiều bất thường cùng lúc, ví dụ vừa có mờ phổi vừa có tràn dịch hoặc tim to.",
        "Trong sản phẩm XAI Medical, hệ thống không sử dụng toàn bộ nhãn để đưa ra báo cáo phức tạp mà trích xuất riêng chỉ số Effusion bằng `model.pathologies.index(\"Effusion\")`. Việc giới hạn phạm vi như vậy giúp sản phẩm tập trung vào một bài toán rõ ràng, giảm nhiễu giao diện và phù hợp với mục tiêu học phần. Tuy nhiên, vì backbone vẫn là mô hình đa nhãn, các đặc trưng học được không chỉ đến từ Effusion mà còn từ kinh nghiệm pretrained trên nhiều bệnh lý lồng ngực khác.",
        "Cách trình bày này cũng giải thích tại sao xác suất thấp không nên hiểu là độ tin cậy thấp của nhãn bình thường. Nếu mô hình trả xác suất Effusion 2%, điều đó có nghĩa xác suất bệnh tràn dịch theo mô hình là thấp, còn độ tin cậy khi hiển thị kết luận bình thường phải được tính là 98% theo logic `1 - probability`. Việc sửa logic hiển thị confidence là một phần quan trọng để tránh gây hiểu nhầm cho người dùng.",
    ],
    "2.4. Fine-tuning classifier Effusion": [
        "Trên nền pretrained DenseNet-121, hệ thống nạp trọng số `pleural_effusion_model_best.pth` để fine-tune classifier phục vụ nhãn Effusion. Đây là cách tiếp cận chuyển giao học hợp lý: mô hình nền đã học nhiều đặc trưng tổng quát của ảnh X-quang từ lượng dữ liệu rất lớn, còn bước fine-tune điều chỉnh phần phân loại cho mục tiêu cụ thể. So với huấn luyện từ đầu, fine-tune tiết kiệm dữ liệu, thời gian và tài nguyên tính toán hơn, đồng thời tận dụng được năng lực biểu diễn đã có.",
        "Cấu hình fine-tune được ghi nhận gồm optimizer AdamW, learning rate 1e-4, 5 epochs, loss function BCELoss và Mixed Precision AMP GradScaler. Việc dùng AdamW giúp kiểm soát weight decay tách biệt với gradient update, thường ổn định trong fine-tuning. Learning rate 1e-4 là mức thận trọng, tránh phá vỡ pretrained features quá mạnh. Số epoch 5 phù hợp với giai đoạn tinh chỉnh nhanh, đặc biệt khi mục tiêu là điều chỉnh classifier chứ không đào tạo một mô hình hoàn toàn mới.",
        "Một chi tiết kỹ thuật quan trọng là Patch v8 chuyển từ BCEWithLogitsLoss sang BCELoss vì đầu ra của model XRV đã qua Sigmoid. Nếu dùng BCEWithLogitsLoss trong khi đầu ra đã Sigmoid, hệ thống có thể gặp lỗi Double Sigmoid, làm biến dạng gradient và xác suất. Việc sửa loss function là ví dụ cho thấy báo cáo không chỉ mô tả lý thuyết mà còn ghi nhận lỗi triển khai thực tế, cách phát hiện và cách khắc phục phù hợp với framework đang dùng.",
    ],
    "2.5. Grad-CAM tại denseblock4": [
        "Grad-CAM cần một tầng feature map sâu để tính bản đồ chú ý. Ban đầu có thể nghĩ đến tầng `norm5`, nhưng trong TorchXRayVision tồn tại nguy cơ lỗi autograd do ReLU inplace liên quan đến tầng này. Khi hook vào một vị trí bị biến đổi inplace, quá trình backward gradient có thể thất bại hoặc trả gradient không hợp lệ. Vì vậy hệ thống chọn `denseblock4`, tầng ngay trước `norm5`, làm target layer cho Grad-CAM. Lựa chọn này an toàn hơn và vẫn giữ đủ thông tin sâu.",
        "`denseblock4` cung cấp 1024 kênh đặc trưng, phản ánh các mẫu hình học và cường độ ảnh mà DenseNet học được ở mức cao. Tuy nhiên, feature map ở tầng sâu thường có độ phân giải không gian thấp, ví dụ khoảng 7 x 7, nên khi phóng to về ảnh gốc, heatmap có thể lan rộng và không trùng khít đường biên tổn thương. Đây là giới hạn tự nhiên của Grad-CAM trong mô hình phân loại, không phải lỗi riêng của sản phẩm. Báo cáo cần trình bày trung thực điểm này để người đọc hiểu đúng ý nghĩa heatmap.",
        "Hậu xử lý heatmap trong sản phẩm gồm alpha blending và contour mảnh để làm kết quả dễ đọc hơn. Hệ thống từng thử lung mask bằng OpenCV nhưng sau đó gỡ bỏ vì thuật toán truyền thống có thể cắt nhầm vùng dịch trắng trong tràn dịch màng phổi. Quyết định gỡ bỏ lung mask thể hiện nguyên tắc an toàn: không thêm thuật toán có thể làm mất thông tin bệnh lý khi chưa có mô hình phân đoạn đáng tin cậy. Vì người dùng yêu cầu không huấn luyện AI mới, báo cáo chỉ ghi nhận U-Net như hướng phát triển tương lai, không đưa vào lõi hiện tại.",
    ],
}


CH3_PARAS = {
    "3.1. Kiến trúc tổng thể": [
        "Kiến trúc hệ thống được tổ chức theo mô hình nhiều lớp, trong đó frontend đảm nhiệm tương tác người dùng, backend xử lý API và điều phối pipeline, AI engine thực hiện suy luận và giải thích, còn SQLite lưu lịch sử phân tích. Cách tách lớp này giúp hệ thống dễ bảo trì, vì giao diện có thể thay đổi mà không làm ảnh hưởng trực tiếp đến mô hình, hoặc mô hình có thể được cập nhật mà không cần viết lại toàn bộ frontend. Đây là nguyên tắc thiết kế quan trọng khi chuyển từ notebook nghiên cứu sang sản phẩm web.",
        "Người dùng thao tác qua giao diện React/Vite, tải ảnh X-quang lên, nhập ghi chú lâm sàng nếu cần và xem kết quả trong modal. Frontend không tự suy luận mà gửi request đến FastAPI, nơi thực hiện kiểm tra file, tiền xử lý, gọi mô hình, tạo heatmap, lưu dữ liệu và trả JSON kết quả. Kết quả gồm nhãn dự đoán, xác suất Effusion, confidence hiển thị, ảnh heatmap, giải thích y khoa và thông tin lưu lịch sử. Cấu trúc API rõ ràng giúp hệ thống có thể mở rộng sang mobile app hoặc dịch vụ khác trong tương lai.",
        "Một điểm đáng chú ý là hệ thống không chỉ xử lý từng ảnh mà còn hỗ trợ batch diagnosis. Trong môi trường thực tế, bác sĩ hoặc người vận hành có thể cần kiểm tra nhiều ảnh cùng lúc, do đó batch API và số worker backend giúp cải thiện hiệu quả. Tuy nhiên, xử lý hàng loạt cũng đặt ra yêu cầu kiểm soát lỗi tốt hơn: nếu một ảnh sai định dạng hoặc ngoài phân phối, hệ thống cần ghi nhận lỗi riêng thay vì làm hỏng toàn bộ batch. Điều này thể hiện tư duy sản phẩm thay vì chỉ tập trung vào demo một ảnh.",
    ],
    "3.2. Backend FastAPI": [
        "Backend sử dụng FastAPI vì framework này nhẹ, nhanh, hỗ trợ typing rõ ràng và phù hợp với các dịch vụ AI Python. Các endpoint chính gồm `/health` để kiểm tra trạng thái, `/stats` để thống kê dữ liệu, `/predict` để phân tích một ảnh, `/predict_batch` để xử lý nhiều ảnh và các endpoint liên quan đến lịch sử. FastAPI đóng vai trò trung gian giữa frontend và mô hình, giúp che giấu độ phức tạp của PyTorch/TorchXRayVision khỏi phía giao diện.",
        "Trong endpoint suy luận, backend thực hiện nhiều bước liên tiếp: nhận file, kiểm tra kích thước, đọc ảnh, chuyển grayscale, normalize, resize, tạo tensor, gọi mô hình, lấy index Effusion, tính xác suất, so sánh threshold, tạo Grad-CAM, hậu xử lý ảnh heatmap và trả kết quả. Mỗi bước đều có nguy cơ lỗi riêng, ví dụ file hỏng, ảnh không phải X-quang, tensor sai shape, model chưa load hoặc Grad-CAM không có gradient. Vì vậy code cần có kiểm tra ngoại lệ và thông báo lỗi rõ ràng thay vì trả lỗi 500 chung chung.",
        "Một lỗi từng xảy ra trong quá trình phát triển là backend chạy sai môi trường Python hoặc không tự reload khi code thay đổi. Việc sửa `run.py` để dùng đúng virtual environment và thêm reload giúp giảm tình trạng frontend gọi vào backend cũ. Đây là bài học triển khai quan trọng: một mô hình đúng trong code chưa chắc đã chạy trong service đang hoạt động nếu process cũ còn giữ cổng 8000 hoặc server chưa reload. Báo cáo ghi nhận yếu tố này để cho thấy sản phẩm đã xử lý vấn đề vận hành thực tế.",
    ],
    "3.3. Frontend React/Vite": [
        "Frontend được thiết kế với các trang chính như Dashboard, Diagnosis, Batch Diagnosis, Compare View, About và Settings. Trang Diagnosis là trung tâm trải nghiệm người dùng, cho phép kéo thả ảnh, hiển thị trạng thái phân tích, mở modal kết quả và xuất PDF. React/Vite phù hợp vì tốc độ phát triển nhanh, hot reload tốt và dễ tách component như ConfidenceGauge, ResultModal, Toast hoặc upload area. Việc tách component giúp sửa lỗi giao diện mà không ảnh hưởng đến logic suy luận.",
        "Một lỗi UX quan trọng đã được phát hiện là cách hiển thị confidence cho ảnh bình thường. Ban đầu hệ thống có thể lấy xác suất Effusion thấp, ví dụ 2%, rồi hiển thị như độ tin cậy của nhãn bình thường, khiến người dùng tưởng AI chỉ tự tin 2%. Logic đúng là nếu kết luận có bệnh thì confidence bằng xác suất bệnh, còn nếu kết luận bình thường thì confidence bằng 100% trừ xác suất bệnh. Cách sửa này không thay đổi mô hình nhưng thay đổi cách truyền đạt kết quả, từ đó giảm hiểu nhầm trong bối cảnh y tế.",
        "Chức năng xuất PDF cũng là phần quan trọng của frontend. Thay vì vẽ từng dòng text thủ công bằng jsPDF, cách mới dùng html2canvas chụp trực tiếp vùng báo cáo rồi đưa vào PDF, nhờ vậy giữ nguyên bố cục tiếng Việt, ảnh gốc, heatmap, màu sắc và spacing. Đây là lựa chọn thực dụng vì báo cáo PDF cần giống giao diện người dùng đã kiểm tra, đặc biệt với tiếng Việt có dấu và layout chứa ảnh. Việc xuất báo cáo giúp sản phẩm không chỉ là công cụ xem tức thời mà còn tạo hồ sơ có thể lưu trữ.",
    ],
    "3.4. Lưu trữ SQLite và lịch sử phân tích": [
        "SQLite được dùng để lưu lịch sử phân tích vì phù hợp với prototype và ứng dụng local/near-production quy mô nhỏ. Cơ sở dữ liệu có thể lưu tên bệnh nhân, ghi chú, đường dẫn ảnh, heatmap, nhãn dự đoán, độ tin cậy và timestamp. Nhờ có lịch sử, người dùng có thể xem lại các lần chẩn đoán, so sánh kết quả và xuất lại báo cáo nếu cần. Đây là điểm khác biệt giữa demo mô hình và một sản phẩm có workflow sử dụng thực tế.",
        "Việc lưu trữ cũng đặt ra yêu cầu bảo mật và riêng tư dữ liệu. Ảnh X-quang và thông tin bệnh nhân là dữ liệu nhạy cảm, do đó trong phiên bản triển khai thực tế cần có xác thực người dùng, phân quyền, mã hóa hoặc ít nhất là kiểm soát thư mục lưu trữ. Trong phạm vi báo cáo học phần, SQLite được mô tả như giải pháp lưu lịch sử để phục vụ demo và nghiên cứu, chưa nên xem là hạ tầng bệnh viện hoàn chỉnh. Đây là giới hạn cần ghi rõ để tránh vượt quá phạm vi sản phẩm hiện tại.",
        "Một lợi ích khác của lưu lịch sử là tạo nền tảng đánh giá dài hạn. Nếu hệ thống lưu lại ảnh, kết quả AI và kết luận của bác sĩ sau này, nhóm phát triển có thể xây dựng bộ dữ liệu kiểm chứng nội bộ, phát hiện lỗi theo nhóm ảnh, phân tích false positive/false negative và cải thiện giao diện. Tuy nhiên, việc dùng dữ liệu lâm sàng thật cần tuân thủ đạo đức nghiên cứu và quy định bảo vệ dữ liệu cá nhân, không thể tự động thu thập hoặc dùng lại nếu chưa được phép.",
    ],
    "3.5. Kiểm soát ảnh ngoài phân phối": [
        "Một hệ thống AI y tế cần tránh trả kết quả tự tin cho dữ liệu không thuộc miền huấn luyện. Nếu người dùng upload ảnh phong cảnh, ảnh tài liệu hoặc ảnh màu không phải X-quang, mô hình DenseNet vẫn có thể trả một xác suất nào đó vì mạng nơ-ron luôn tính toán trên tensor đầu vào. Tuy nhiên xác suất đó không có ý nghĩa y khoa. Vì vậy module OODDetector được dùng để kiểm tra kích thước, tỷ lệ khung hình, entropy và đặc điểm màu/độ sáng trước khi đưa ảnh vào pipeline suy luận.",
        "OOD không giải quyết toàn bộ rủi ro, nhưng là lớp bảo vệ cần thiết. Ảnh X-quang thật cũng có thể có chất lượng kém, bị xoay, bị crop, có nhãn chữ, có thiết bị y tế hoặc khác biệt so với dữ liệu huấn luyện. Do đó kiểm soát OOD nên được xem là bước lọc ban đầu, không phải bảo đảm tuyệt đối. Trong phiên bản tương lai, có thể bổ sung kiểm tra DICOM metadata, phân loại view PA/AP, phát hiện vùng phổi hoặc kiểm tra chất lượng ảnh theo tiêu chuẩn chẩn đoán hình ảnh.",
        "Trong báo cáo, OOD được đặt ở vị trí trước preprocessing và suy luận để nhấn mạnh nguyên tắc an toàn. Hệ thống chỉ nên trả kết quả khi ảnh đầu vào đủ điều kiện tối thiểu, còn nếu ảnh không phù hợp thì trả thông báo giải thích cho người dùng. Cách làm này tốt hơn việc cố gắng dự đoán mọi ảnh, vì trong y tế, một lời từ chối hợp lý thường an toàn hơn một kết quả sai có vẻ chắc chắn.",
    ],
}


CH4_PARAS = {
    "4.1. Kết quả đánh giá mô hình": [
        "Theo tài liệu dự án Patch v8, mô hình đạt Accuracy 81,46% và AUC-ROC 0,87 trên hold-out test set trong quá trình fine-tuning Cloud/Kaggle. Accuracy cho biết tỷ lệ dự đoán đúng tổng thể tại ngưỡng vận hành, còn AUC phản ánh khả năng phân biệt lớp trên nhiều ngưỡng khác nhau. Trong bài toán y tế, AUC thường có ý nghĩa quan trọng vì nó cho thấy mô hình có khả năng xếp hạng ca bệnh cao hơn ca bình thường hay không, ngay cả khi ngưỡng cuối cùng có thể điều chỉnh theo mục tiêu lâm sàng.",
        "Ngưỡng phân loại hiện tại là 0,0682, tức 6,82%, được chọn bằng Youden's J statistic trên đường cong ROC. Ngưỡng này thấp hơn nhiều so với 0,5 mặc định, nhưng điều đó có thể hợp lý trong bối cảnh sàng lọc y tế, nơi bỏ sót bệnh thường nguy hiểm hơn cảnh báo sai. Tuy nhiên, lựa chọn ngưỡng thấp cũng có thể làm tăng false positive, do đó khi triển khai thực tế cần đánh giá thêm Precision, Recall, Specificity và confusion matrix trên tập test đại diện.",
        "Báo cáo không tự đưa ra TP/FP/TN/FN vì hiện chưa có file export cứng cho ma trận nhầm lẫn trong source code cục bộ. Đây là điểm cần minh bạch: thay vì bịa số liệu để báo cáo đẹp hơn, tài liệu ghi rõ cần bổ sung artifact thực nghiệm. Cách trình bày này phù hợp với chuẩn học thuật, vì số liệu định lượng phải có nguồn đo rõ ràng, có tập test xác định và có khả năng kiểm tra lại.",
    ],
    "4.2. Đánh giá heatmap và giải thích": [
        "Heatmap Grad-CAM giúp người dùng quan sát vùng ảnh ảnh hưởng đến quyết định của mô hình. Trong tràn dịch màng phổi, vùng đáy phổi và góc sườn hoành thường là khu vực quan trọng vì dịch có xu hướng tích tụ ở vị trí thấp theo trọng lực. Nếu heatmap tập trung vào các vùng này, người dùng có thêm cơ sở để hiểu dự đoán Effusion. Tuy nhiên, heatmap cũng có thể lan ra ngoài hoặc tập trung vào vùng không trực quan do mô hình phân loại học shortcut hoặc do feature map có độ phân giải thấp.",
        "Trong quá trình phát triển, hệ thống từng gặp vấn đề contour quá dày, vẽ nhiều vùng nhiễu và heatmap bị biến dạng tỷ lệ. Các lỗi này được xử lý bằng cách giữ aspect ratio khi overlay, làm mỏng contour, lọc vùng nóng chính và dùng alpha blending để chỉ tô màu vùng có nhiệt thay vì phủ toàn bộ ảnh bằng nền màu. Những cải tiến này không thay đổi mô hình, nhưng cải thiện đáng kể khả năng đọc kết quả. Đây là ví dụ điển hình cho sự khác biệt giữa chất lượng mô hình và chất lượng sản phẩm.",
        "Hệ thống cũng từng thử lung mask bằng OpenCV để ép heatmap nằm trong vùng phổi, nhưng thuật toán này có thể cắt nhầm vùng dịch vì tràn dịch biểu hiện trắng đục, trong khi phương pháp threshold truyền thống thường tìm vùng khí tối. Do người dùng yêu cầu giữ nguyên mô hình đã huấn luyện và không thêm AI mới, giải pháp an toàn là gỡ bỏ lung mask truyền thống, giữ heatmap gốc và chỉ dùng hậu xử lý trực quan nhẹ. Quyết định này tránh làm mất vùng bệnh quan trọng chỉ vì một thuật toán ảnh cổ điển nhận sai.",
    ],
    "4.3. Đánh giá hệ thống phần mềm": [
        "Về mặt phần mềm, hệ thống đã đạt các chức năng chính của một prototype hoàn chỉnh: upload ảnh, phân tích đơn lẻ, phân tích hàng loạt, hiển thị dashboard, lưu lịch sử, xem lại kết quả, tạo heatmap, sinh giải thích và xuất báo cáo PDF. Backend FastAPI giúp API rõ ràng, frontend React giúp giao diện linh hoạt, SQLite giúp triển khai gọn trong môi trường local. Những thành phần này tạo thành một workflow sử dụng được, không chỉ là notebook nghiên cứu.",
        "Một số lỗi triển khai đã được phát hiện và xử lý trong quá trình phát triển, ví dụ lỗi backend import sai môi trường Python, lỗi process cũ giữ cổng 8000, lỗi click upload hai lần do input file không reset, lỗi confidence hiển thị sai khi ảnh bình thường, lỗi PDF font/bố cục tiếng Việt và lỗi Grad-CAM khi hook vào tầng không phù hợp. Việc ghi nhận các lỗi này làm báo cáo thực tế hơn, vì một sản phẩm AI không chỉ phụ thuộc vào mô hình mà còn phụ thuộc vào rất nhiều chi tiết vận hành và UX.",
        "Tuy nhiên, hệ thống vẫn còn các điểm cần hoàn thiện trước khi triển khai trong môi trường nghiêm túc. Cần có test tự động cho backend, frontend và pipeline AI; cần chuẩn hóa logging; cần xử lý bảo mật dữ liệu bệnh nhân; cần hỗ trợ DICOM; cần đóng gói Docker; cần cấu hình GPU/CPU rõ ràng; và cần có quy trình kiểm thử với dữ liệu y tế được gán nhãn bởi chuyên gia. Những yêu cầu này vượt quá phạm vi báo cáo học phần nhưng là hướng đi cần thiết nếu muốn phát triển sản phẩm xa hơn.",
    ],
    "4.4. So sánh với hướng mở rộng U-Net và ViT": [
        "U-Net và Attention U-Net thường được nhắc đến trong bài toán phân đoạn ảnh y tế vì có khả năng tạo mask vùng giải phẫu hoặc tổn thương. Nếu trong tương lai hệ thống cần định lượng diện tích tràn dịch hoặc giới hạn heatmap trong phổi một cách chính xác, mô hình phân đoạn có thể hữu ích. Tuy nhiên, trong phiên bản hiện tại, hệ thống không dùng U-Net và cũng không huấn luyện AI thứ hai. Báo cáo chỉ nêu U-Net như hướng phát triển, không mô tả nó như thành phần đã có trong sản phẩm.",
        "Vision Transformer cũng là hướng nghiên cứu đáng quan tâm vì có khả năng mô hình hóa quan hệ toàn cục giữa các vùng ảnh. Đối với X-quang ngực, quan hệ giữa đáy phổi, tim, cơ hoành và khoang màng phổi có thể quan trọng. Tuy vậy, ViT thường cần dữ liệu lớn và cấu hình huấn luyện cẩn thận. Trong phạm vi sản phẩm hiện tại, DenseNet-121 vẫn là lựa chọn thực dụng hơn vì đã có pretrained weights chuyên miền X-quang, pipeline TorchXRayVision ổn định và khả năng triển khai nhẹ hơn.",
        "Việc so sánh các hướng mở rộng giúp báo cáo có chiều sâu nhưng vẫn không làm sai hiện trạng sản phẩm. Điểm cốt lõi là phân biệt giữa thành phần đã triển khai và thành phần đề xuất. Thành phần đã triển khai gồm DenseNet-121, fine-tuned Effusion classifier, Grad-CAM, FastAPI, React, SQLite và PDF export. Thành phần đề xuất gồm U-Net, Attention U-Net, ViT, DICOM đầy đủ, Docker, authentication và bộ artifact đánh giá mở rộng. Cách trình bày này tránh việc báo cáo thổi phồng năng lực hệ thống.",
    ],
    "4.5. Rủi ro, giới hạn và đạo đức sử dụng": [
        "Rủi ro lớn nhất của hệ thống AI y tế là người dùng hiểu kết quả như kết luận tuyệt đối. Mặc dù mô hình có AUC 0,87 và Accuracy 81,46%, vẫn có khả năng sai, đặc biệt với ảnh chất lượng thấp, bệnh lý phối hợp, dữ liệu khác phân phối huấn luyện hoặc ca lâm sàng hiếm. Vì vậy giao diện và báo cáo PDF cần ghi rõ hệ thống chỉ hỗ trợ tham khảo, không thay thế bác sĩ. Cảnh báo này không làm giảm giá trị sản phẩm mà giúp sản phẩm được sử dụng đúng vai trò.",
        "Một giới hạn khác là thiếu confusion matrix và các chỉ số theo lớp được export cứng. Accuracy và AUC là quan trọng nhưng chưa đủ để đánh giá lâm sàng. Ví dụ một mô hình có accuracy tốt vẫn có thể bỏ sót ca bệnh nếu dữ liệu mất cân bằng; một mô hình có recall cao có thể tạo nhiều false positive nếu threshold quá thấp. Do đó phiên bản tiếp theo cần lưu đầy đủ TP, FP, TN, FN, Precision, Recall, F1-score, Specificity, ROC curve và PR curve để đánh giá toàn diện hơn.",
        "Về đạo đức dữ liệu, ảnh X-quang và thông tin bệnh nhân phải được xem là dữ liệu nhạy cảm. Nếu hệ thống được triển khai ngoài môi trường học tập, cần có quy trình ẩn danh dữ liệu, xin phép sử dụng, kiểm soát truy cập, mã hóa lưu trữ và chính sách xóa dữ liệu. Với dữ liệu công khai như NIH, CheXpert, MIMIC-CXR, PadChest hoặc covid-chestxray-dataset, nhóm phát triển cần tuân thủ giấy phép và điều kiện sử dụng. Báo cáo vì vậy không chỉ mô tả kỹ thuật mà còn đặt sản phẩm trong trách nhiệm sử dụng AI y tế.",
    ],
}

EXTRA_CH1 = {
    "1.6. Đặc điểm bệnh học của tràn dịch màng phổi trên X-quang": [
        "Tràn dịch màng phổi là biểu hiện hình ảnh của việc dịch tích tụ trong khoang màng phổi, làm thay đổi độ cản quang và hình thái vùng đáy phổi. Trên X-quang ngực thẳng, dấu hiệu thường gặp là tù góc sườn hoành, mờ đồng nhất vùng đáy phổi hoặc đường cong giới hạn dịch tùy lượng dịch. Trong các trường hợp lượng dịch nhỏ, dấu hiệu có thể rất kín đáo và dễ bị bỏ qua nếu chất lượng ảnh thấp hoặc bệnh nhân chụp không đúng tư thế. Vì vậy bài toán phát hiện Effusion phù hợp với mô hình sàng lọc, bởi mô hình có thể học các tín hiệu thống kê từ nhiều ảnh và đưa ra cảnh báo để người đọc phim kiểm tra kỹ hơn.",
        "Tuy nhiên, biểu hiện hình ảnh của Effusion không phải lúc nào cũng đơn giản. Một vùng trắng ở đáy phổi có thể liên quan đến dịch, xẹp phổi, viêm phổi, bóng tim, tư thế chụp hoặc chất lượng tia. Ngược lại, một trường hợp có dịch ít có thể chỉ làm mờ nhẹ góc sườn hoành mà không tạo vùng trắng rõ ràng. Điều này giải thích vì sao mô hình phân loại chỉ nên được xem là công cụ hỗ trợ xác suất, còn đánh giá cuối cùng cần kết hợp với bối cảnh lâm sàng, phim nghiêng, siêu âm hoặc các phương tiện chẩn đoán khác nếu cần.",
        "Khi đưa vào báo cáo học phần, phần bệnh học giúp người đọc hiểu vì sao hệ thống không thể chỉ dựa trên thuật toán vẽ vùng nóng. Heatmap có thể chỉ ra nơi mô hình chú ý, nhưng diễn giải y khoa cần dựa trên kiến thức về vị trí giải phẫu, tư thế bệnh nhân và dấu hiệu hình ảnh. Do đó, phần giải thích trong sản phẩm nên dùng ngôn ngữ thận trọng như “vùng ảnh có ảnh hưởng đến dự đoán” thay vì khẳng định “đây là vùng bệnh chính xác”."
    ],
    "1.7. Khác biệt giữa mô hình phân loại và mô hình phân đoạn": [
        "Một điểm thường gây nhầm lẫn trong XAI y tế là đánh đồng heatmap của mô hình phân loại với mask phân đoạn. Mô hình phân loại học cách dự đoán nhãn toàn ảnh, vì vậy nó không được huấn luyện bằng đường biên tổn thương và cũng không có mục tiêu tối ưu hóa độ chính xác pixel. Ngược lại, mô hình phân đoạn như U-Net cần dữ liệu mask do chuyên gia gán nhãn, từ đó học cách dự đoán từng pixel thuộc vùng tổn thương hay không. Hai loại mô hình phục vụ hai mục tiêu khác nhau và không nên đánh giá bằng cùng một tiêu chuẩn.",
        "Trong sản phẩm XAI Medical hiện tại, DenseNet-121 là mô hình phân loại đa nhãn. Nó trả xác suất Effusion cho toàn ảnh, sau đó Grad-CAM truy vết vùng ảnh có ảnh hưởng đến xác suất đó. Nếu người dùng kỳ vọng heatmap giống đường biên tổn thương chính xác, họ sẽ dễ đánh giá sai công nghệ. Vì vậy giao diện và báo cáo cần giải thích rằng heatmap là công cụ diễn giải quyết định mô hình, không phải kết quả phân đoạn được xác nhận bởi chuyên gia.",
        "Sự phân biệt này cũng ảnh hưởng đến hướng phát triển. Nếu mục tiêu tương lai là định lượng diện tích dịch hoặc vẽ ranh giới phổi chính xác, cần dữ liệu mask và mô hình phân đoạn. Nếu mục tiêu là sàng lọc nhanh ca nghi ngờ, mô hình phân loại kèm Grad-CAM là lựa chọn nhẹ hơn, dễ triển khai hơn và phù hợp với nguồn lực hiện có. Báo cáo vì vậy chọn cách mô tả đúng bản chất sản phẩm thay vì thổi phồng heatmap thành khả năng phân đoạn."
    ],
    "1.8. Yêu cầu minh bạch trong hệ thống AI y tế": [
        "Minh bạch trong AI y tế không chỉ là việc tạo heatmap. Một hệ thống minh bạch cần cho biết mô hình dùng dữ liệu nào, nhãn nào, preprocessing ra sao, ngưỡng quyết định bao nhiêu, chỉ số đánh giá được đo trên tập nào và còn thiếu artifact nào. Nếu chỉ trình bày Accuracy mà không nói tập test, nếu chỉ đưa heatmap mà không giải thích giới hạn, hoặc nếu hiển thị confidence sai logic, hệ thống có thể gây hiểu nhầm dù mô hình bên trong tương đối tốt. Do đó báo cáo này dành nhiều dung lượng cho nguồn dữ liệu, threshold, fine-tuning và giới hạn sử dụng.",
        "Trong sản phẩm hiện tại, minh bạch được thể hiện ở nhiều tầng: tầng dữ liệu ghi rõ pretrained weights đến từ bốn dataset lớn; tầng mô hình ghi rõ DenseNet-121, Effusion index và classifier fine-tuned; tầng giải thích dùng Grad-CAM tại denseblock4; tầng giao diện hiển thị confidence theo nhãn dự đoán; tầng báo cáo PDF lưu lại ảnh, kết quả và diễn giải. Các tầng này bổ sung cho nhau, giúp người dùng không phải nhìn mô hình như một hộp đen hoàn toàn.",
        "Tuy vậy, minh bạch cũng cần đi kèm khiêm tốn khoa học. Báo cáo ghi rõ chưa có TP/FP/TN/FN export cứng và chưa có confusion matrix trong source cục bộ. Việc thừa nhận thiếu sót này là cần thiết vì báo cáo học thuật phải dựa trên bằng chứng. Một hệ thống có thể được đánh giá là nghiêm túc hơn khi nó nói rõ điều chưa biết, thay vì cố tạo ra các con số không có nguồn đo."
    ],
    "1.9. Định hướng sản phẩm so với bài toán nghiên cứu": [
        "Một bài toán nghiên cứu thường tập trung vào mô hình, thuật toán và chỉ số đánh giá, trong khi một sản phẩm cần thêm trải nghiệm người dùng, vận hành, lưu trữ, xử lý lỗi và xuất báo cáo. XAI Medical nằm giữa hai hướng này: phần lõi dựa trên mô hình học sâu và XAI, nhưng phần triển khai hướng đến một ứng dụng web sử dụng được. Vì vậy báo cáo cần cân bằng giữa phân tích thuật toán và mô tả kỹ thuật phần mềm, tránh chỉ trình bày như một notebook huấn luyện mô hình.",
        "Việc có frontend, backend, database và PDF export làm tăng giá trị thực tiễn của đề tài. Người dùng không cần chạy script Python mà có thể thao tác qua giao diện; kết quả không chỉ in ra console mà được hiển thị bằng ảnh, số liệu và giải thích; lịch sử không mất sau mỗi lần chạy mà được lưu trong SQLite; báo cáo có thể tải về để phục vụ trình bày hoặc lưu hồ sơ. Những chi tiết này phản ánh tư duy xây dựng sản phẩm, vốn là yêu cầu quan trọng trong các đề tài ứng dụng.",
        "Tuy nhiên, định hướng sản phẩm cũng kéo theo trách nhiệm lớn hơn. Khi một mô hình được đóng gói thành giao diện đẹp, người dùng có thể tin tưởng nó hơn mức cần thiết. Do đó báo cáo luôn gắn phần sản phẩm với cảnh báo giới hạn y tế, yêu cầu kiểm thử và nguyên tắc không thay thế bác sĩ. Đây là cách tiếp cận an toàn khi đưa AI từ môi trường thử nghiệm sang bối cảnh có liên quan đến sức khỏe."
    ],
}

EXTRA_CH2 = {
    "2.6. Tổ chức tập train, validation và test": [
        "Trong kịch bản fine-tuning Cloud/Kaggle, dữ liệu được chia theo tỷ lệ 70% Train, 10% Validation và 20% Test. Tập train dùng để cập nhật trọng số classifier, tập validation dùng để theo dõi quá trình học và lựa chọn cấu hình tốt hơn, còn tập test giữ độc lập để đánh giá cuối cùng. Cách chia này là thực hành phổ biến trong học máy vì giúp giảm nguy cơ đánh giá lạc quan trên chính dữ liệu đã dùng để điều chỉnh mô hình.",
        "Điều quan trọng là tập test phải được giữ kín trong quá trình lựa chọn mô hình. Nếu liên tục xem kết quả test để chỉnh threshold, chọn epoch hoặc sửa pipeline, tập test sẽ dần trở thành validation trá hình và chỉ số cuối cùng không còn phản ánh khả năng tổng quát. Vì vậy khi báo cáo Accuracy 81,46% và AUC 0,87, cần ghi rõ đây là kết quả trên hold-out test set trong quá trình fine-tuning Cloud/Kaggle. Bộ 50 ảnh COVID-19/Pneumonia dùng để test hành vi sản phẩm không thay thế được tập test chuẩn.",
        "Nếu có điều kiện hoàn thiện báo cáo, nhóm phát triển nên lưu lại file mô tả split, seed random, danh sách ảnh trong từng tập và script đánh giá. Những artifact này giúp tái lập kết quả, chứng minh rằng dữ liệu không bị rò rỉ giữa train và test, đồng thời hỗ trợ kiểm tra lại khi có thay đổi pipeline. Đây là yêu cầu quan trọng nếu sản phẩm được nâng cấp từ báo cáo học phần lên nghiên cứu nghiêm túc."
    ],
    "2.7. Vai trò của loss function và lỗi Double Sigmoid": [
        "Loss function quyết định cách mô hình bị phạt khi dự đoán sai, do đó có ảnh hưởng trực tiếp đến quá trình fine-tune. Trong bài toán nhị phân hoặc đa nhãn, hai lựa chọn thường gặp là BCELoss và BCEWithLogitsLoss. BCEWithLogitsLoss bao gồm bước Sigmoid nội bộ, phù hợp khi model trả logits chưa qua Sigmoid. Ngược lại, BCELoss phù hợp khi đầu ra đã là xác suất trong khoảng [0, 1]. Nếu dùng sai, gradient và xác suất có thể bị xử lý hai lần hoặc không đúng miền.",
        "Trong dự án này, Patch v8 ghi nhận việc chuyển sang BCELoss vì đầu ra của model XRV đã có Sigmoid. Đây là bản sửa quan trọng để tránh lỗi Double Sigmoid. Lỗi này có thể khiến xác suất bị nén, mô hình học kém hoặc chỉ số đánh giá không phản ánh đúng năng lực thực. Việc phát hiện và sửa loss function cho thấy quá trình xây dựng hệ thống không chỉ là lắp ghép thư viện mà cần hiểu rõ giả định toán học của từng thành phần.",
        "Khi đưa nội dung này vào báo cáo, cần trình bày bằng ngôn ngữ dễ hiểu nhưng chính xác: nếu mô hình đã tự biến logits thành xác suất, ta không nên dùng loss tiếp tục áp dụng Sigmoid. Cách giải thích này giúp người đọc hiểu vì sao một thay đổi nhỏ trong code huấn luyện lại có thể ảnh hưởng lớn đến kết quả. Đây cũng là phần thể hiện năng lực kỹ thuật của nhóm khi xử lý lỗi sâu trong pipeline học máy."
    ],
    "2.8. Ý nghĩa của ngưỡng 0,0682 trong y tế": [
        "Ngưỡng 0,0682 có vẻ rất thấp nếu so với trực giác thông thường 50%, nhưng trong bài toán y tế, ngưỡng tối ưu không nhất thiết là 0,5. Nếu mục tiêu là sàng lọc bệnh, hệ thống thường ưu tiên phát hiện nhiều ca nghi ngờ hơn, chấp nhận một số cảnh báo sai để giảm nguy cơ bỏ sót. Youden's J statistic là phương pháp chọn ngưỡng dựa trên đường cong ROC nhằm cân bằng độ nhạy và độ đặc hiệu. Vì vậy ngưỡng 0,0682 cần được hiểu trong bối cảnh tối ưu hóa lâm sàng, không phải do mô hình yếu.",
        "Tuy nhiên, ngưỡng thấp cũng có mặt trái. Nếu dữ liệu triển khai khác nhiều so với dữ liệu test, hệ thống có thể tạo nhiều false positive, làm tăng tải kiểm tra lại cho bác sĩ. Do đó khi chuyển sản phẩm sang môi trường mới, ngưỡng cần được hiệu chỉnh lại trên dữ liệu nội bộ có nhãn chuẩn. Báo cáo nên coi 0,0682 là ngưỡng hiện tại của phiên bản v8, không phải giá trị bất biến cho mọi bệnh viện hoặc mọi thiết bị chụp.",
        "Trong giao diện, threshold nên được tách khỏi confidence. Threshold dùng để quyết định nhãn; confidence dùng để truyền đạt mức độ chắc chắn theo nhãn đã hiển thị. Nếu xác suất Effusion là 0,02 và threshold là 0,0682, nhãn là bình thường trong phạm vi Effusion, còn confidence hiển thị nên là 98%. Cách trình bày này giúp người dùng hiểu đúng ý nghĩa xác suất và tránh lỗi UX từng xảy ra trong hệ thống."
    ],
    "2.9. Kiểm soát chất lượng dữ liệu đầu vào": [
        "Chất lượng dữ liệu đầu vào ảnh hưởng trực tiếp đến chất lượng suy luận. Một ảnh bị xoay, crop quá mạnh, chụp thiếu vùng đáy phổi hoặc có độ tương phản kém có thể làm mô hình trả xác suất không ổn định. Vì mô hình được huấn luyện trên ảnh X-quang ngực theo một số chuẩn nhất định, ảnh quá khác biệt so với phân phối huấn luyện cần được cảnh báo. Đây là lý do hệ thống có kiểm tra ảnh ngoài phân phối trước khi suy luận.",
        "Kiểm soát chất lượng không chỉ là bài toán kỹ thuật mà còn là yêu cầu nghiệp vụ. Nếu người dùng upload ảnh không phù hợp mà hệ thống vẫn trả “bình thường” hoặc “tràn dịch”, kết quả đó có thể tạo cảm giác an toàn giả. Một thông báo từ chối như “ảnh không giống X-quang ngực hợp lệ” tuy có vẻ ít ấn tượng hơn dự đoán, nhưng an toàn và trung thực hơn. Với hệ thống y tế, từ chối dự đoán khi dữ liệu không đủ điều kiện là hành vi đúng.",
        "Trong tương lai, kiểm soát chất lượng có thể được mở rộng bằng các tiêu chí chuyên sâu như phát hiện view PA/AP, kiểm tra vùng phổi có đủ trong ảnh, phát hiện ảnh quá sáng/quá tối, kiểm tra DICOM metadata hoặc dùng mô hình phân loại chất lượng ảnh. Tuy nhiên, các hướng này cần dữ liệu và đánh giá riêng, không nên gộp vào năng lực hiện tại nếu chưa triển khai. Báo cáo vì vậy chỉ mô tả OOD hiện tại ở mức kiểm tra cơ bản."
    ],
}

EXTRA_CH3 = {
    "3.6. Thiết kế API và dữ liệu trả về": [
        "API của hệ thống cần trả về dữ liệu đủ cho frontend hiển thị nhưng không quá phụ thuộc vào cách trình bày giao diện. Một response tốt thường gồm nhãn dự đoán, xác suất Effusion, confidence hiển thị, đường dẫn hoặc base64 heatmap, thông điệp giải thích, trạng thái OOD và mã lỗi nếu có. Khi frontend nhận cấu trúc rõ ràng, các component như ResultModal, ConfidenceGauge hoặc PDF export có thể hoạt động ổn định mà không phải suy luận lại logic nghiệp vụ.",
        "Thiết kế API cũng cần phân biệt lỗi hệ thống và lỗi dữ liệu người dùng. Nếu model chưa load hoặc Grad-CAM thiếu gradient, đó là lỗi backend cần log và sửa. Nếu người dùng upload ảnh không phải X-quang, đó là lỗi đầu vào cần trả thông báo dễ hiểu. Phân biệt này giúp hệ thống thân thiện hơn và dễ debug hơn. Trong quá trình phát triển, lỗi “Backend responded with an error” quá chung khiến việc tìm nguyên nhân khó hơn, nên báo cáo khuyến nghị chuẩn hóa error response.",
        "Với batch diagnosis, API cần trả kết quả theo từng ảnh thay vì chỉ trả một trạng thái chung. Một batch có thể chứa cả ảnh hợp lệ và ảnh lỗi; nếu toàn bộ batch thất bại vì một ảnh sai, trải nghiệm người dùng sẽ kém. Thiết kế tốt là mỗi item có trạng thái riêng, thông báo riêng và kết quả riêng. Điều này đặc biệt quan trọng nếu hệ thống dùng trong môi trường có nhiều ảnh cần sàng lọc nhanh."
    ],
    "3.7. Thiết kế giao diện kết quả": [
        "Giao diện kết quả phải trình bày thông tin theo thứ tự ưu tiên: kết luận chính, độ tin cậy, ảnh gốc, heatmap, giải thích y khoa và hành động tiếp theo. Nếu đưa quá nhiều số liệu lên đầu, người dùng không chuyên có thể khó nắm ý chính; nếu chỉ đưa kết luận mà không có heatmap, hệ thống thiếu tính giải thích. Result Modal trong sản phẩm đóng vai trò gom các thông tin này vào một cửa sổ tập trung, giúp người dùng xem nhanh nhưng vẫn có đủ bằng chứng trực quan.",
        "Màu sắc trong giao diện y tế cần dùng thận trọng. Màu đỏ/cam có thể biểu thị vùng nghi ngờ hoặc cảnh báo, màu xanh/cyan có thể biểu thị trạng thái an toàn hoặc confidence, nhưng nếu dùng quá nhiều sẽ gây nhiễu. Heatmap nên nổi bật vùng có ý nghĩa nhưng không che mất cấu trúc X-quang bên dưới. Vì vậy alpha blending là lựa chọn phù hợp: nơi không có nhiệt giữ ảnh gốc, nơi có nhiệt mới được tô màu, giúp người đọc vẫn thấy cấu trúc giải phẫu.",
        "Ngoài hiển thị chính, giao diện cần có các hành động rõ ràng như “chẩn đoán tiếp theo”, “xem lại”, “xuất báo cáo” và thông báo toast khi hoàn tất. Những chi tiết nhỏ này làm sản phẩm dễ sử dụng hơn. Trong báo cáo học phần, việc mô tả UX không phải phần phụ, vì một hệ thống AI chỉ có giá trị khi người dùng có thể hiểu và thao tác đúng với kết quả của nó."
    ],
    "3.8. Quản lý file ảnh và heatmap": [
        "Khi người dùng upload ảnh, hệ thống cần quyết định cách lưu file gốc, file heatmap và thông tin liên quan. Lưu trực tiếp ảnh trong database có thể làm database phình to, trong khi lưu đường dẫn file giúp nhẹ hơn nhưng cần quản lý thư mục và quyền truy cập. Trong prototype, lưu đường dẫn ảnh và heatmap là lựa chọn đơn giản. Tuy nhiên, khi triển khai thực tế, cần có chiến lược đặt tên file, tránh trùng lặp, tránh lộ thông tin bệnh nhân và dọn dữ liệu cũ.",
        "Heatmap cũng cần được lưu cùng ngữ cảnh mô hình. Nếu sau này thay model hoặc threshold, cùng một ảnh có thể cho heatmap khác. Vì vậy mỗi bản ghi lịch sử nên lưu phiên bản hệ thống, ví dụ Patch v8, threshold 0,0682, target layer denseblock4 và weights sử dụng. Điều này giúp truy vết kết quả, đặc biệt khi người dùng mở lại báo cáo cũ. Một hệ thống y tế nghiêm túc cần khả năng audit, không chỉ lưu ảnh cuối cùng.",
        "Quản lý file còn liên quan đến chức năng PDF. Nếu PDF chỉ chứa đường dẫn ảnh local, khi chuyển máy có thể mất hình. Cách xuất hiện tại dùng render vùng báo cáo thành ảnh trong PDF, giúp báo cáo tự chứa nội dung hiển thị. Đây là ưu điểm thực tế, nhất là khi người dùng cần gửi file cho người khác hoặc nộp báo cáo học phần mà không kèm thư mục ảnh."
    ],
    "3.9. Khả năng mở rộng và triển khai": [
        "Trong môi trường phát triển, hệ thống có thể chạy bằng `python run.py`, frontend dev server và SQLite local. Tuy nhiên, nếu muốn triển khai ổn định hơn, cần đóng gói backend và frontend bằng Docker, cấu hình biến môi trường, quản lý model weights, mount thư mục lưu ảnh và thiết lập reverse proxy. Docker giúp giảm lỗi khác môi trường, ví dụ global Python không có `torchxrayvision` trong khi virtual environment của dự án có đủ dependency.",
        "Khả năng mở rộng cũng liên quan đến tài nguyên phần cứng. DenseNet-121 có thể chạy CPU cho demo nhưng GPU sẽ giúp batch diagnosis nhanh hơn. Nếu triển khai nhiều người dùng, backend cần kiểm soát hàng đợi, giới hạn kích thước file, timeout và số worker để tránh quá tải. Đối với ảnh y tế, không nên chỉ tối ưu tốc độ mà bỏ qua kiểm tra chất lượng, vì một phản hồi nhanh nhưng sai hoặc không kiểm soát lỗi có thể nguy hiểm hơn phản hồi chậm nhưng chắc chắn.",
        "Báo cáo định hướng triển khai theo từng giai đoạn: giai đoạn học phần tập trung prototype local; giai đoạn tiếp theo chuẩn hóa Docker, logging, test tự động; giai đoạn sau mới tính đến bảo mật, phân quyền, DICOM và tích hợp hệ thống bệnh viện. Cách chia giai đoạn giúp đề tài thực tế hơn, tránh đặt mục tiêu quá lớn so với phạm vi hiện tại."
    ],
}

EXTRA_CH4 = {
    "4.6. Quy trình kiểm thử đề xuất": [
        "Quy trình kiểm thử nên bắt đầu từ unit test cho các hàm preprocessing, ví dụ kiểm tra ảnh màu có được chuyển grayscale đúng không, ảnh sau resize có kích thước 224 x 224 không và normalize có đưa về miền giá trị phù hợp không. Tiếp theo là test API với ảnh hợp lệ, ảnh sai định dạng, ảnh quá lớn và ảnh ngoài phân phối. Các test này giúp phát hiện lỗi kỹ thuật trước khi đánh giá mô hình, vì một pipeline sai có thể làm mọi chỉ số mô hình trở nên vô nghĩa.",
        "Ở tầng mô hình, cần test consistency bằng một số ảnh cố định và snapshot kết quả xác suất trong khoảng chấp nhận được. Nếu sau khi sửa code, cùng một ảnh cho xác suất lệch quá mạnh, nhóm phát triển cần kiểm tra xem preprocessing, weights hoặc threshold có thay đổi không. Với Grad-CAM, có thể kiểm tra file heatmap được tạo, kích thước đúng, không rỗng toàn số 0 và không gây lỗi gradient. Đây là các kiểm tra tự động thực dụng, không thay thế đánh giá y khoa nhưng giúp bảo vệ pipeline.",
        "Ở tầng giao diện, cần test các thao tác upload, xem modal, xuất PDF, xem lịch sử và batch diagnosis. Các lỗi UX như phải click hai lần để chọn lại file hoặc confidence hiển thị sai cần được đưa vào test hồi quy nếu có điều kiện. Một hệ thống AI y tế tốt không chỉ đúng ở model mà còn phải ổn định ở toàn bộ đường đi từ ảnh đầu vào đến báo cáo đầu ra."
    ],
    "4.7. Phân tích lỗi thường gặp": [
        "Lỗi đầu tiên thường gặp là lỗi môi trường chạy. Nếu backend được khởi động bằng Python global thiếu `torchxrayvision`, hệ thống sẽ báo module not found dù virtual environment dự án đã cài đủ. Cách khắc phục là `run.py` phải dùng đúng interpreter trong môi trường dự án. Đây là lỗi vận hành phổ biến khi dự án AI có nhiều dependency nặng và nhiều môi trường Python cùng tồn tại trên máy.",
        "Lỗi thứ hai là lỗi process cũ giữ cổng hoặc không reload code mới. Khi người dùng sửa `xai_gradcam.py` nhưng backend cũ vẫn chạy, frontend tiếp tục gọi vào code cũ khiến kết quả heatmap không đổi. Điều này dễ gây hiểu nhầm rằng thuật toán sửa không hiệu quả. Bài học là cần log version, restart rõ ràng và có endpoint health trả thông tin phiên bản model/code đang chạy. Khi hệ thống có version rõ, người dùng dễ biết mình đang test đúng bản hay chưa.",
        "Lỗi thứ ba là lỗi diễn giải kết quả. Một xác suất Effusion thấp không có nghĩa AI thiếu tự tin vào kết luận bình thường; nó có nghĩa xác suất bệnh thấp. Nếu giao diện hiển thị sai, người dùng có thể đánh giá sai mô hình. Đây là loại lỗi không nằm trong model nhưng ảnh hưởng trực tiếp đến niềm tin người dùng. Vì vậy kiểm thử sản phẩm cần bao gồm cả logic hiển thị, không chỉ kiểm thử hàm predict."
    ],
    "4.8. Tiêu chí nghiệm thu sản phẩm": [
        "Một bộ tiêu chí nghiệm thu hợp lý nên gồm bốn nhóm: chức năng, kỹ thuật, dữ liệu và trình bày. Về chức năng, hệ thống phải upload được ảnh, trả kết quả, tạo heatmap, lưu lịch sử, xử lý batch và xuất PDF. Về kỹ thuật, backend không lỗi 500 với ảnh hợp lệ, OOD hoạt động với ảnh không phù hợp, Grad-CAM không rỗng và PDF giữ được tiếng Việt. Về dữ liệu, báo cáo phải ghi rõ nguồn pretrained, tập fine-tune, threshold và thiếu sót artifact. Về trình bày, giao diện và báo cáo phải dễ đọc.",
        "Đối với mô hình, nghiệm thu không nên chỉ dựa vào một vài ảnh demo. Cần có tập test có nhãn, thống kê số ca Effusion và bình thường, lưu kết quả dự đoán và tính các chỉ số. Nếu có thể, nên nhờ người có chuyên môn xem lại các ca sai để phân loại nguyên nhân: ảnh chất lượng thấp, bệnh lý khác, mô hình bỏ sót, threshold chưa phù hợp hoặc lỗi preprocessing. Phân tích lỗi như vậy có giá trị hơn việc chỉ báo cáo một con số accuracy.",
        "Đối với báo cáo học phần, tiêu chí nghiệm thu còn bao gồm quy cách trình bày: đúng font Times New Roman 13, giãn dòng 1.5, đánh số chương/mục, có bảng từ viết tắt, mục lục, mục lục hình/bảng, tài liệu tham khảo và nội dung nằm trong khoảng 30 đến 100 trang. File được tạo trong lượt này bám theo các yêu cầu đó, đồng thời giữ văn phong đoạn văn dài theo yêu cầu của người dùng."
    ],
    "4.9. Đề xuất cải tiến sau báo cáo": [
        "Cải tiến ưu tiên đầu tiên là xuất đầy đủ artifact đánh giá. Nhóm phát triển nên tạo script đánh giá cố định đọc test set, chạy model, lưu CSV gồm tên ảnh, nhãn thật, xác suất, nhãn dự đoán và threshold, sau đó sinh confusion matrix, ROC curve, PR curve và bảng chỉ số. Khi có artifact này, báo cáo sau sẽ có bằng chứng định lượng mạnh hơn, tránh phụ thuộc vào mô tả thủ công. Đây là bước cần thiết nếu muốn bảo vệ kết quả trước hội đồng hoặc dùng làm nền cho bài nghiên cứu.",
        "Cải tiến thứ hai là chuẩn hóa dữ liệu đầu vào và hỗ trợ DICOM. Trong thực tế bệnh viện, ảnh X-quang thường ở định dạng DICOM với metadata quan trọng về thiết bị, tư thế, windowing và thông tin kỹ thuật. Hỗ trợ DICOM giúp hệ thống gần thực tế hơn, nhưng cũng yêu cầu xử lý bảo mật và ẩn danh tốt hơn. Đây là hướng phát triển lớn, nên được tách thành giai đoạn riêng thay vì gộp vào prototype hiện tại.",
        "Cải tiến thứ ba là bổ sung cơ chế phản hồi của bác sĩ. Nếu bác sĩ có thể xác nhận kết quả đúng/sai hoặc ghi chú lý do, hệ thống sẽ tạo được dữ liệu đánh giá nội bộ theo thời gian. Tuy nhiên, phản hồi này phải được quản lý cẩn thận và không tự động dùng để huấn luyện lại nếu chưa có quy trình kiểm duyệt. Một vòng phản hồi tốt có thể biến sản phẩm từ công cụ demo thành nền tảng hỗ trợ nghiên cứu lâm sàng có kiểm soát."
    ],
}


def add_sections(doc: Document, sections: dict[str, list[str]]):
    for title, paras in sections.items():
        add_h2(doc, title)
        for p in paras:
            add_para(doc, p)


def add_references(doc: Document):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run("DANH MỤC TÀI LIỆU THAM KHẢO")
    r.bold = True
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(14)
    refs = [
        "G. Huang, Z. Liu, L. van der Maaten, K. Q. Weinberger (2016), “Densely Connected Convolutional Networks,” arXiv:1608.06993.",
        "X. Wang et al. (2017), “ChestX-ray8: Hospital-scale Chest X-ray Database and Benchmarks on Weakly-Supervised Classification and Localization of Common Thorax Diseases,” CVPR.",
        "J. Irvin et al. (2019), “CheXpert: A Large Chest Radiograph Dataset with Uncertainty Labels and Expert Comparison,” AAAI.",
        "A. E. W. Johnson et al. (2019), “MIMIC-CXR, a de-identified publicly available database of chest radiographs with free-text reports,” Scientific Data.",
        "A. Bustos et al. (2020), “PadChest: A large chest x-ray image dataset with multi-label annotated reports,” Medical Image Analysis.",
        "J. P. Cohen et al. (2021), “TorchXRayVision: A library of chest X-ray datasets and models,” arXiv:2111.00595.",
        "R. R. Selvaraju et al. (2016), “Grad-CAM: Visual Explanations from Deep Networks via Gradient-based Localization,” arXiv:1610.02391.",
        "A. Chattopadhyay et al. (2017), “Grad-CAM++: Improved Visual Explanations for Deep Convolutional Networks,” arXiv:1710.11063.",
        "O. Ronneberger, P. Fischer, T. Brox (2015), “U-Net: Convolutional Networks for Biomedical Image Segmentation,” arXiv:1505.04597.",
        "J. P. Cohen et al., “COVID-19 Image Data Collection,” GitHub repository ieee8023/covid-chestxray-dataset.",
        "Tài liệu nội bộ dự án XAI Medical: mã nguồn backend/frontend, Patch v8, cấu hình mô hình và báo cáo kiểm thử trong repository D:\\My\\XAI.",
    ]
    for i, ref in enumerate(refs, 1):
        add_para(doc, f"[{i}] {ref}")


def main():
    images = create_images()
    doc = Document()
    set_document_defaults(doc)
    cover(doc)
    front_matter(doc)

    add_chapter(doc, "CHƯƠNG 1. TỔNG QUAN VÀ CƠ SỞ LÝ THUYẾT")
    add_sections(doc, {**SECTION_PARAS, **EXTRA_CH1})
    add_picture(doc, images[0], "Hình 1.1. Kiến trúc tổng thể hệ thống XAI Medical.")

    add_chapter(doc, "CHƯƠNG 2. DỮ LIỆU, MÔ HÌNH HUẤN LUYỆN VÀ PHƯƠNG PHÁP")
    add_sections(doc, {**CH2_PARAS, **EXTRA_CH2})
    add_picture(doc, images[1], "Hình 2.1. Quy mô các nguồn dữ liệu pretrained trong weights all.")
    add_table(doc, ["Nguồn dữ liệu", "Quy mô xấp xỉ", "Vai trò"], [
        ["NIH ChestX-ray14", "112.120 ảnh", "Nguồn dữ liệu X-quang ngực quy mô lớn, gắn với nhãn bệnh lý lồng ngực."],
        ["CheXpert", "224.316 ảnh", "Bổ sung ảnh X-quang với nhãn không chắc chắn và quy mô bệnh viện lớn."],
        ["MIMIC-CXR", "377.110 ảnh", "Nguồn dữ liệu lâm sàng lớn, có ảnh và báo cáo đã khử định danh."],
        ["PadChest", "~160.000 ảnh", "Bổ sung đa dạng nguồn ảnh và nhãn bệnh lý X-quang."],
    ], [3.5, 3.0, 8.0])
    add_caption(doc, "Bảng 2.1. Nguồn dữ liệu và quy mô pretrained của mô hình gốc.")
    add_table(doc, ["Nhóm", "Nội dung"], [
        ["Backbone", "DenseNet-121 từ TorchXRayVision, weights densenet121-res224-all."],
        ["Fine-tune", "pleural_effusion_model_best.pth cho classifier Effusion."],
        ["Hyperparameters", "AdamW, learning rate 1e-4, 5 epochs, BCELoss, AMP GradScaler."],
        ["Preprocessing", "Grayscale, XRayResizer(224), xrv.datasets.normalize(img_np, 255)."],
        ["Threshold", "0,0682 theo Youden's J statistic trên ROC."],
    ], [4.0, 10.5])
    add_caption(doc, "Bảng 2.2. Tóm tắt cấu hình dữ liệu, mô hình và huấn luyện.")
    add_picture(doc, images[2], "Hình 2.2. Pipeline tiền xử lý ảnh và suy luận Effusion.")
    add_picture(doc, images[3], "Hình 2.3. Nguyên lý Grad-CAM tại denseblock4.")

    add_chapter(doc, "CHƯƠNG 3. THIẾT KẾ VÀ XÂY DỰNG HỆ THỐNG")
    add_sections(doc, {**CH3_PARAS, **EXTRA_CH3})
    add_table(doc, ["Thành phần", "Công nghệ", "Chức năng chính"], [
        ["Frontend", "React/Vite", "Upload ảnh, hiển thị kết quả, dashboard, batch diagnosis, xuất PDF."],
        ["Backend", "FastAPI/Uvicorn", "Nhận request, kiểm tra dữ liệu, gọi mô hình, trả JSON và lưu lịch sử."],
        ["AI Engine", "PyTorch/TorchXRayVision", "DenseNet-121, Effusion probability, Grad-CAM tại denseblock4."],
        ["Database", "SQLite", "Lưu bệnh án, ảnh, heatmap, nhãn dự đoán, độ tin cậy và timestamp."],
        ["PDF", "html2canvas/jsPDF", "Render vùng báo cáo thành PDF giữ bố cục tiếng Việt và hình ảnh."],
    ], [3.0, 3.5, 8.0])
    add_caption(doc, "Bảng 3.1. Thành phần kiến trúc phần mềm của hệ thống.")
    add_picture(doc, images[5], "Hình 3.1. Luồng xuất báo cáo PDF từ giao diện kết quả.")

    add_chapter(doc, "CHƯƠNG 4. THỰC NGHIỆM, ĐÁNH GIÁ VÀ HOÀN THIỆN SẢN PHẨM")
    add_sections(doc, {**CH4_PARAS, **EXTRA_CH4})
    add_picture(doc, images[4], "Hình 4.1. Minh họa ngưỡng vận hành theo ROC và Youden's J statistic.")
    add_table(doc, ["Chỉ số/Thành phần", "Giá trị hiện có", "Ý nghĩa"], [
        ["Accuracy", "81,46%", "Tỷ lệ dự đoán đúng trên hold-out test set Cloud/Kaggle."],
        ["AUC-ROC", "0,87", "Khả năng phân biệt lớp trên nhiều ngưỡng."],
        ["Optimal threshold", "0,0682", "Ngưỡng Effusion chọn theo Youden's J statistic."],
        ["Input size", "224 x 224", "Kích thước tensor đầu vào cho DenseNet-121."],
        ["Missing artifact", "Chưa export TP/FP/TN/FN", "Cần bổ sung confusion matrix để nghiệm thu định lượng sâu."],
    ], [3.5, 3.5, 7.5])
    add_caption(doc, "Bảng 4.1. Thông số đánh giá và triển khai đã xác nhận.")

    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run("KẾT LUẬN")
    r.bold = True
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(14)
    add_para(doc, "Báo cáo đã trình bày đầy đủ quá trình xây dựng hệ thống XAI Medical theo bốn nhóm nội dung: cơ sở lý thuyết, dữ liệu và mô hình huấn luyện, thiết kế hệ thống và đánh giá sản phẩm. Kết quả cho thấy sản phẩm đã kết hợp được backbone DenseNet-121 pretrained trên nhiều dataset X-quang lớn, bước fine-tune classifier Effusion, pipeline preprocessing đúng chuẩn TorchXRayVision, Grad-CAM tại denseblock4 và hệ thống web hoàn chỉnh gồm FastAPI, React, SQLite và xuất PDF. Điểm mạnh của đề tài là không dừng ở mô hình mà xây dựng được workflow sản phẩm có thể thao tác, lưu lịch sử và tạo báo cáo.")
    add_para(doc, "Nhược điểm chính hiện tại là báo cáo chưa có artifact thực nghiệm đầy đủ như confusion matrix, Precision, Recall, F1-score, Specificity và ROC/PR curve được export thành file kiểm chứng. Heatmap Grad-CAM cũng cần được hiểu đúng như bản đồ chú ý của mô hình phân loại, không phải mask phân đoạn vùng bệnh. Ngoài ra, hệ thống vẫn cần hoàn thiện bảo mật dữ liệu, hỗ trợ DICOM, test tự động, Docker và quy trình kiểm thử trên dữ liệu có chuyên gia xác nhận nếu muốn tiến gần hơn đến môi trường ứng dụng thực tế.")
    add_para(doc, "Hướng phát triển tiếp theo là chuẩn hóa bộ đánh giá, bổ sung báo cáo huấn luyện có thể tái lập, cải thiện OOD, tối ưu giao diện cho bác sĩ, mở rộng quản lý bệnh án và cân nhắc mô hình phân đoạn trong tương lai khi có dữ liệu/nguồn lực phù hợp. Trong phạm vi hiện tại, sản phẩm đã đáp ứng mục tiêu học phần: thể hiện được khả năng ứng dụng AI giải thích được trong bài toán y tế, đồng thời trình bày rõ nguồn dữ liệu, mô hình, phương pháp và giới hạn sử dụng.")
    add_references(doc)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()

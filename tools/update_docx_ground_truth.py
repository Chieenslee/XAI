from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt


SRC = Path(r"D:\My\XAI\Doc\Bao_cao_XAI_Medical_Tieng_Viet.docx")
OUT = Path(r"D:\My\XAI\Doc\Bao_cao_XAI_Medical_Tieng_Viet_GroundTruth.docx")
BACKUP = Path(r"D:\My\XAI\Doc\Bao_cao_XAI_Medical_Tieng_Viet.before_ground_truth_update.docx")


PARAGRAPH_REPLACEMENTS = {
    3: (
        "Báo cáo này trình bày sản phẩm XAI Medical, một hệ thống hỗ trợ chẩn đoán tràn dịch màng phổi "
        "từ ảnh X-quang ngực. Hệ thống sử dụng backbone DenseNet-121 của TorchXRayVision với weights "
        "`densenet121-res224-all`, vốn được huấn luyện trên tổ hợp NIH ChestX-ray14, CheXpert, MIMIC-CXR "
        "và PadChest. Trên nền pretrained này, hệ thống nạp trọng số fine-tuned "
        "`models/pleural_effusion_model_best.pth` để tập trung vào nhãn Effusion, kết hợp Grad-CAM tại "
        "`denseblock4`, FastAPI, React/Vite, SQLite và chức năng xuất báo cáo PDF."
    ),
    4: (
        "Từ khóa: XAI, tràn dịch màng phổi, X-quang ngực, DenseNet-121, TorchXRayVision, "
        "densenet121-res224-all, Effusion, Grad-CAM, FastAPI, React."
    ),
    17: (
        "Các nghiên cứu về chẩn đoán ảnh X-quang bằng học sâu thường đi theo ba hướng: phân loại bệnh lý, "
        "định vị/giải thích vùng tổn thương và phân đoạn cấu trúc giải phẫu. Các bộ dữ liệu lớn như "
        "NIH ChestX-ray14, CheXpert, MIMIC-CXR và PadChest tạo nền tảng cho mô hình pretrained đa nhãn. "
        "Trong sản phẩm này, mô hình không học riêng một bệnh từ đầu mà kế thừa weights `all` của "
        "TorchXRayVision rồi fine-tune classifier cho Effusion."
    ),
    18: (
        "TorchXRayVision cung cấp pipeline chuẩn hóa cho ảnh X-quang ngực và danh sách pathology đa nhãn. "
        "Hệ thống trích xuất riêng chỉ số `Effusion` bằng `model.pathologies.index(\"Effusion\")`. Với bài toán "
        "giải thích, Grad-CAM được dùng để chỉ ra vùng ảnh ảnh hưởng đến quyết định của mô hình; cần hiểu đây "
        "là bản đồ chú ý của mô hình phân loại, không phải mask phân đoạn vùng bệnh."
    ),
    21: (
        "Mô hình nền sử dụng `xrv.models.DenseNet(weights=\"densenet121-res224-all\")`. Tên weights `all` cho biết "
        "mô hình gốc được huấn luyện trên tổ hợp bốn bộ dữ liệu X-quang lớn: NIH ChestX-ray14, CheXpert, "
        "MIMIC-CXR và PadChest, với tổng quy mô hơn 800.000 ảnh. Trong đó NIH ChestX-ray14 đóng góp 112.120 ảnh, "
        "CheXpert khoảng 224.316 ảnh, MIMIC-CXR khoảng 377.110 ảnh và PadChest khoảng 160.000 ảnh."
    ),
    22: "Đọc ảnh từ file upload JPEG/PNG, kiểm tra kích thước và từ chối ảnh ngoài phân phối không giống ảnh X-quang.",
    23: (
        "Bắt buộc chuyển ảnh về grayscale; ảnh màu phong cảnh hoặc ảnh không phù hợp miền X-quang sẽ bị hệ thống "
        "kiểm tra và từ chối."
    ),
    24: "Resize ảnh về 224 x 224 bằng `XRayResizer(224)` để phù hợp DenseNet-121 pretrained của TorchXRayVision.",
    25: (
        "Chuẩn hóa bằng `xrv.datasets.normalize(img_np, 255)`, đưa pixel từ miền [0, 255] về chuẩn X-quang "
        "[-1024, 1024], không dùng mean/std ImageNet."
    ),
    26: (
        "Tạo tensor đầu vào kích thước (1, 1, 224, 224), sau đó lấy xác suất lớp Effusion từ danh sách pathology "
        "đa nhãn của mô hình."
    ),
    27: (
        "Trong kịch bản fine-tuning trên Cloud/Kaggle, chia dữ liệu theo tỷ lệ 70% Train, 10% Validation và "
        "20% Test. Bộ 50 ảnh COVID-19/Pneumonia tải từ repo `ieee8023/covid-chestxray-dataset` chỉ dùng để "
        "kiểm thử thực tế, không phải tập test tạo ra Accuracy 81,46% và AUC 0,87."
    ),
    30: (
        "Backbone chính là DenseNet-121 từ TorchXRayVision với weights `densenet121-res224-all`. Đây là mô hình "
        "phân loại đa nhãn cho nhiều bệnh lý lồng ngực như Atelectasis, Consolidation, Infiltration, "
        "Pneumothorax, Edema, Emphysema, Fibrosis, Effusion, Pneumonia, Pleural Thickening, Cardiomegaly, "
        "Nodule, Mass và Hernia. Hệ thống chỉ trích xuất nhãn Effusion để phục vụ bài toán tràn dịch màng phổi."
    ),
    31: "B. Fine-tuning và giải thích bằng Grad-CAM",
    32: (
        "Sau khi lấy pretrained backbone, hệ thống nạp trọng số `pleural_effusion_model_best.pth` để fine-tune "
        "classifier cho Effusion. Cấu hình fine-tune ghi nhận gồm AdamW, learning rate 1e-4, 5 epochs, BCELoss "
        "và Mixed Precision AMP GradScaler. Việc dùng BCELoss thay cho BCEWithLogitsLoss là bản vá v8 nhằm tránh "
        "lỗi Double Sigmoid, vì đầu ra TorchXRayVision đã qua Sigmoid. Module XAI hook vào `denseblock4` thay vì "
        "`norm5` để tránh lỗi autograd do ReLU inplace, đồng thời giữ 1024 kênh đặc trưng sâu cho Grad-CAM."
    ),
    36: (
        "Theo tài liệu dự án Patch v8, mô hình đạt Accuracy 81,46% và AUC-ROC 0,87 trên hold-out test set trong "
        "quá trình fine-tuning Cloud/Kaggle. Đây không phải kết quả đo trên bộ 50 ảnh COVID-19/Pneumonia dùng "
        "để kiểm thử thực tế. Ngưỡng vận hành hiện tại là 0,0682, được chọn bằng Youden's J statistic trên "
        "đường cong ROC để cân bằng True Positive và False Positive, thay vì dùng ngưỡng 0,5 mặc định."
    ),
    38: (
        "Hiện chưa có file export cứng cho ma trận nhầm lẫn TP/FP/TN/FN trong source code cục bộ. Do đó báo cáo "
        "nghiệm thu học thuật cần bổ sung confusion matrix, Precision, Recall, F1-score, độ đặc hiệu, ROC artifact "
        "và mô tả rõ tập test trước khi kết luận định lượng sâu hơn."
    ),
    50: (
        "XAI Medical là một sản phẩm prototype/near-production cho bài toán hỗ trợ chẩn đoán tràn dịch màng phổi. "
        "Điểm mạnh của hệ thống là sử dụng backbone DenseNet-121 pretrained trên nhiều dataset X-quang lớn, "
        "fine-tune riêng classifier Effusion, giữ nguyên mô hình đã huấn luyện và chỉ dùng hậu xử lý để cải thiện "
        "hiển thị heatmap, báo cáo và trải nghiệm người dùng."
    ),
    51: (
        "Các hướng phát triển tiếp theo gồm: xuất đầy đủ confusion matrix và ROC artifact, chuẩn hóa báo cáo "
        "huấn luyện, bổ sung test tự động, triển khai Docker, nâng cấp OOD, thêm xác thực người dùng, hỗ trợ DICOM "
        "và cải thiện báo cáo PDF. Không nên bổ sung U-Net/AI thứ hai vào mô tả hiện tại nếu chưa có quyết định "
        "huấn luyện hoặc tích hợp mô hình mới."
    ),
}


REFERENCE_ADDITIONS = [
    "[10] X. Wang et al., “ChestX-ray8: Hospital-scale Chest X-ray Database and Benchmarks on Weakly-Supervised Classification and Localization of Common Thorax Diseases,” CVPR, 2017.",
    "[11] J. Irvin et al., “CheXpert: A Large Chest Radiograph Dataset with Uncertainty Labels and Expert Comparison,” AAAI, 2019.",
    "[12] A. E. W. Johnson et al., “MIMIC-CXR, a de-identified publicly available database of chest radiographs with free-text reports,” Scientific Data, 2019.",
    "[13] A. Bustos et al., “PadChest: A large chest x-ray image dataset with multi-label annotated reports,” Medical Image Analysis, 2020.",
    "[14] J. P. Cohen et al., “COVID-19 Image Data Collection,” GitHub repository ieee8023/covid-chestxray-dataset.",
]


def set_paragraph_text(paragraph, text: str) -> None:
    style = paragraph.style
    alignment = paragraph.alignment
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.add_run(text)
    paragraph.style = style
    paragraph.alignment = alignment
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)


def set_cell(cell, text: str) -> None:
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)


def add_table_row(table, values: list[str]) -> None:
    row = table.add_row()
    for cell, value in zip(row.cells, values):
        set_cell(cell, value)


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    if not BACKUP.exists():
        shutil.copy2(SRC, BACKUP)

    doc = Document(SRC)

    for idx, text in PARAGRAPH_REPLACEMENTS.items():
        set_paragraph_text(doc.paragraphs[idx], text)

    # Cover metadata table.
    set_cell(doc.tables[0].cell(2, 1), "DenseNet-121/TorchXRayVision weights densenet121-res224-all, Grad-CAM, FastAPI, React Vite, SQLite")

    # Related work table.
    set_cell(doc.tables[2].cell(1, 2), "Backbone chính: DenseNet-121 pretrained weights densenet121-res224-all.")
    set_cell(doc.tables[2].cell(2, 0), "Grad-CAM")
    set_cell(doc.tables[2].cell(2, 2), "Hook tại denseblock4 để tạo heatmap giải thích cho nhãn Effusion.")
    set_cell(doc.tables[2].cell(3, 2), "Không dùng trong phiên bản hiện tại; chỉ nên nêu như hướng mở rộng nếu có mô hình mới.")

    # Pipeline table.
    set_cell(doc.tables[3].cell(2, 2), "Nhận file, kiểm tra định dạng/kích thước, kiểm tra OOD và gọi pipeline grayscale + normalize.")
    set_cell(doc.tables[3].cell(3, 2), "Sinh probability cho lớp Effusion từ mô hình đa nhãn TorchXRayVision.")
    set_cell(doc.tables[3].cell(5, 1), "Grad-CAM")
    set_cell(doc.tables[3].cell(5, 2), "Tạo heatmap từ denseblock4; contour/alpha blending chỉ là hậu xử lý hiển thị, không huấn luyện AI mới.")

    # Evaluation table.
    set_cell(doc.tables[4].cell(2, 1), "0,87")
    set_cell(doc.tables[4].cell(2, 2), "Khả năng phân biệt lớp trên hold-out test set Cloud/Kaggle.")
    set_cell(doc.tables[4].cell(3, 2), "Ngưỡng Effusion theo Youden's J statistic trên ROC.")
    add_table_row(doc.tables[4], ["Training datasets", "NIH ChestX-ray14, CheXpert, MIMIC-CXR, PadChest", "Nguồn pretrained weights densenet121-res224-all."])
    add_table_row(doc.tables[4], ["Training scale", ">800.000 ảnh", "Tổng quy mô xấp xỉ từ bốn bộ dữ liệu X-quang lớn."])
    add_table_row(doc.tables[4], ["Fine-tune weights", "pleural_effusion_model_best.pth", "Trọng số classifier Effusion nạp tại thư mục models/."])
    add_table_row(doc.tables[4], ["Fine-tune config", "AdamW, LR 1e-4, 5 epochs, BCELoss, AMP", "BCELoss dùng để tránh Double Sigmoid trong Patch v8."])
    add_table_row(doc.tables[4], ["Test caveat", "Chưa export TP/FP/TN/FN", "Cần bổ sung confusion matrix và ROC artifact khi nghiệm thu học thuật."])

    # References.
    doc.paragraphs[61].add_run().add_break(WD_BREAK.LINE)
    for ref in REFERENCE_ADDITIONS:
        p = doc.add_paragraph(ref)
        p.style = doc.paragraphs[53].style
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

    doc.save(OUT)
    print(f"Updated: {OUT}")
    print(f"Backup: {BACKUP}")


if __name__ == "__main__":
    main()
